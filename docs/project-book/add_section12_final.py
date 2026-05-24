"""
add_section12_final.py
Adds the final professional summary section (Section 12) to RECIPEBOX_PROJECT_BOOK.docx
and applies consistency improvements to existing headings/paragraphs where possible.
Written at Mahat academic Hebrew standard.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = "docs/project-book/RECIPEBOX_PROJECT_BOOK.docx"

doc = Document(DOC_PATH)

# ─────────────────────────────────────────────
# Helper functions (consistent with all previous sections)
# ─────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p


def add_subheading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return p


def add_sub3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════
# PAGE BREAK — start Section 12 on a new page
# ═══════════════════════════════════════════════
doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 12 — סיכום
# ═══════════════════════════════════════════════
add_heading(doc, "12. סיכום", level=1)

add_paragraph(doc,
    "פרויקט RecipeBox הוא יישום רשת מלא (Full-Stack Web Application) "
    "שפותח כפרויקט גמר במוסד להשכלה טכנולוגית. "
    "המערכת מאפשרת לאנשים לאסוף, לנהל ולשתף מתכונים בצורה מאורגנת, "
    "תוך שילוב יכולות בינה מלאכותית לשיפור חוויית השימוש. "
    "הפרק הנוכחי מרכז את עיקרי הפרויקט, מציג את הישגיו המרכזיים "
    "ומשקף על האתגרים שנפגשו במהלך פיתוחו."
)

# ─────────────────────────────────────────────
# 12.1 עיקרי הפרויקט
# ─────────────────────────────────────────────
add_subheading(doc, "12.1 עיקרי הפרויקט")

add_paragraph(doc,
    "RecipeBox פותח כמערכת Client-Server מלאה הכוללת שלוש שכבות: "
    "ממשק משתמש (React + TypeScript), שרת API (ASP.NET Core 9 / C#) "
    "ובסיס נתונים (PostgreSQL). "
    "המערכת תומכת בשלוש רמות משתמשים — אורח, משתמש רגיל ומנהל — "
    "עם הפרדת הרשאות ברורה המיושמת הן בצד הלקוח והן בצד השרת."
)

add_paragraph(doc,
    "להלן הרכיבים המרכזיים שמומשו בפרויקט:"
)

add_bullet(doc,
    "מערכת אימות והרשאות מלאה מבוססת JWT עם תוקף של שבעה ימים, "
    "חיפוש סיסמאות באמצעות BCrypt, וניהול תפקידים (User / Admin).",
    bold_prefix="אימות:"
)
add_bullet(doc,
    "ניהול תוכן (CRUD מלא) למתכונים ואלבומים, "
    "עם תמיכה בשני סוגי מתכון: מתכון מלא (טקסט) ומתכון-קישור (URL חיצוני).",
    bold_prefix="ניהול תוכן:"
)
add_bullet(doc,
    "אחסון תמונות בענן באמצעות Cloudinary CDN, "
    "עם אפשרות העלאת קובץ ישיר או הדבקת כתובת URL.",
    bold_prefix="תמונות:"
)
add_bullet(doc,
    "אינטגרציה עם מודל שפה גדול (Groq / llama-3.3-70b-versatile) "
    "דרך פרוקסי מאובטח בשרת, המאפשרת חיפוש חכם, "
    "שיפור טקסט, הצעת שם ותיאור, המרת יחידות מידה וכפל כמויות.",
    bold_prefix="בינה מלאכותית:"
)
add_bullet(doc,
    "ממשק ניהול (Admin Panel) הכולל סטטיסטיקות מערכת, "
    "ניהול משתמשים (חסימה/ביטול חסימה), ניהול אלבומים ומחיקת מתכונים.",
    bold_prefix="לוח ניהול:"
)
add_bullet(doc,
    "מנגנון מועדפים המבוסס על טבלת קשר (Many-to-Many) בין משתמשים למתכונים, "
    "עם עדכון מצב מקומי ללא טעינה מחדש של הדף.",
    bold_prefix="מועדפים:"
)
add_bullet(doc,
    "חיפוש וסינון מתכונים בצד השרת לפי שם, תיאור ואלבום, "
    "עם מיון אינטליגנטי שמציב מתכוני המשתמש עצמו ראשונים.",
    bold_prefix="חיפוש:"
)

# ─────────────────────────────────────────────
# 12.2 הישגים טכניים
# ─────────────────────────────────────────────
add_subheading(doc, "12.2 הישגים טכניים מרכזיים")

add_paragraph(doc,
    "בתהליך הפיתוח הושגו מספר יעדים טכניים שלא היו ברורים מראש:"
)

achievements = [
    ("הפרדת הרשאות דו-שכבתית",
     "כל הגנה מיושמת פעמיים: פעם אחת בלקוח (ProtectedRoute, הסתרת כפתורים) "
     "ופעם שנייה בשרת ([Authorize], בדיקת בעלות). "
     "גישה זו מבטיחה שגם אם הלקוח נפרץ, השרת ימשיך לאכוף את ההרשאות."),
    ("מחיקה מדורגת (Cascade Delete)",
     "כל הקשרים במסד הנתונים מוגדרים עם ON DELETE CASCADE. "
     "מחיקת אלבום מוחקת אוטומטית את כל מתכוניו ואת רשומות המועדפים שלהם, "
     "מבלי לדרוש קוד נוסף ברמת היישום."),
    ("פרוקסי AI בשרת",
     "גישה לשירותי AI חיצוניים עוברת דרך AIController בשרת בלבד. "
     "כך מפתחות ה-API נשמרים בסביבת השרת ואינם נחשפים לדפדפן."),
    ("Context API ללא Redux",
     "ניהול המצב הגלובלי (אימות, חלון Modal) מומש באמצעות React Context API בלבד, "
     "ללא תלות בספריית Redux — פתרון מינימלי ויעיל לגודל הפרויקט."),
    ("עדכון מצב מקומי (Optimistic UI)",
     "פעולות כמו הוספה/הסרת מועדפים ומחיקת פריטים מעדכנות את ה-State המקומי "
     "מיידית, ללא טעינה מחדש של הנתונים מהשרת, מה שמשפר את תחושת המהירות."),
    ("התאמת ממשק לפי תפקיד",
     "ה-Navbar, כרטיסיות המתכון ודפי הניהול מציגים אפשרויות שונות "
     "בהתאם לתפקיד המשתמש — אורח, משתמש רגיל או מנהל — "
     "מבלי לדרוש קומפוננטות נפרדות לכל תפקיד."),
]

make_table(doc, ["הישג", "תיאור"], achievements)

# ─────────────────────────────────────────────
# 12.3 סקירת ארכיטקטורה סופית
# ─────────────────────────────────────────────
add_subheading(doc, "12.3 סקירת הארכיטקטורה הסופית")

add_paragraph(doc,
    "המערכת שפותחה מציגה ארכיטקטורה בשלוש שכבות עם הפרדת אחריות ברורה:"
)

arch_summary = [
    ("שכבת הלקוח",
     "React 18 + TypeScript",
     "ממשק משתמש, ניהול מצב (Context API), ניתוב (React Router), "
     "אינטגרציה עם Chakra UI. מארגן 13 מסכים ב-SPA."),
    ("שכבת השרת",
     "ASP.NET Core 9 / C#",
     "REST API עם 24 נקודות קצה. ניהול הרשאות (JWT + [Authorize]), "
     "לוגיקה עסקית (Services), פרוקסי ל-Groq AI ו-Cloudinary."),
    ("שכבת הנתונים",
     "PostgreSQL + EF Core",
     "ארבע טבלאות: Users, Albums, Recipes, UserFavorites. "
     "ניהול מיגרציות, אינדקסים וקשרי Cascade."),
    ("שירותים חיצוניים",
     "Cloudinary + Groq AI",
     "Cloudinary: אחסון תמונות ב-CDN. "
     "Groq: מודל שפה גדול לבינה מלאכותית גנרטיבית."),
]

make_table(doc, ["שכבה", "טכנולוגיה", "אחריות"], arch_summary)

# ─────────────────────────────────────────────
# 12.4 מיפוי פרקי הספר
# ─────────────────────────────────────────────
add_subheading(doc, "12.4 מיפוי פרקי ספר הפרויקט")

add_paragraph(doc,
    "ספר הפרויקט בנוי כך שכל פרק מספק שכבת הבנה נוספת על המערכת:"
)

chapters = [
    ("1–5",   "רקע ותכנון",      "הגדרת המטרות, ניתוח שוק, דרישות ותכנון ראשוני"),
    ("6",     "תיאור תפקודי",   "זרימת השימוש, תהליכי ליבה (הרשמה, יצירה, עריכה, מחיקה)"),
    ("7",     "מבנה נתונים",    "ארבע הטבלאות, שדות, מפתחות, קשרים ומיגרציות"),
    ("8",     "לוגיקה מרכזית",  "API, אימות, הרשאות, CRUD, ניהול מצב, AI, טיפול בשגיאות"),
    ("9",     "ארכיטקטורה",     "שלוש שכבות המערכת, מבנה התיקיות, זרימת תקשורת"),
    ("10",    "מדריך משתמש",    "הרשאות לפי תפקיד, זרימות עבודה, מצבי גבול"),
    ("11",    "תיאור מסכים",    "13 המסכים: נתיב, קומפוננטה, הרשאה, קלט, כפתורים, פעולות"),
    ("12",    "סיכום",          "עיקרי הפרויקט, הישגים, לקחים והמלצות להמשך"),
]

make_table(doc, ["פרק", "נושא", "תוכן עיקרי"], chapters)

# ─────────────────────────────────────────────
# 12.5 אתגרים ופתרונות
# ─────────────────────────────────────────────
add_subheading(doc, "12.5 אתגרים ופתרונות")

add_paragraph(doc,
    "במהלך פיתוח הפרויקט נתגלו מספר אתגרים טכניים שדרשו פתרונות מחושבים:"
)

challenges = [
    ("ניהול הרשאות רב-שכבתי",
     "המערכת דורשת בדיקות הרשאה בשתי נקודות שונות — לקוח ושרת — "
     "מה שיכל להוביל לכפילות קוד. "
     "הפתרון: ProtectedRoute בלקוח לניתוב, "
     "[Authorize] ו-Ownership Check בשרת לאכיפה."),
    ("עקביות נתונים בעת מחיקה",
     "מחיקת אלבום יכלה להותיר מתכונים 'יתומים'. "
     "הפתרון: CASCADE DELETE בכל קשרי ה-FK, המבטיח עקביות אוטומטית."),
    ("חשיפת מפתחות API",
     "שירותי AI חיצוניים דורשים מפתח API שלא ניתן לשלוח ללקוח. "
     "הפתרון: AIController משמש כפרוקסי — הלקוח אינו רואה את המפתח לעולם."),
    ("ניהול session בסביבת SPA",
     "רענון הדפדפן מאפס את React State. "
     "הפתרון: שמירת פרטי המשתמש ב-localStorage, "
     "עם Interceptor ב-Axios לניקוי אוטומטי בעת תפוגת הטוקן."),
    ("שמירת מצב ממשק בעת שינויים",
     "עדכון מועדפים גרם לטעינה מחדש של הדף. "
     "הפתרון: עדכון State מקומי ישיר לאחר קבלת תגובת השרת, ללא קריאת API נוספת."),
]

make_table(doc, ["אתגר", "תיאור הפתרון"], challenges)

# ─────────────────────────────────────────────
# 12.6 תרומת הבינה המלאכותית
# ─────────────────────────────────────────────
add_subheading(doc, "12.6 תרומת הבינה המלאכותית למערכת")

add_paragraph(doc,
    "אחת מנקודות הייחוד של RecipeBox היא שילוב יכולות AI ישירות בתוך זרימת העבודה של המשתמש. "
    "הבינה המלאכותית אינה תוספת שולית אלא חלק אורגני מהמוצר:"
)

add_bullet(doc,
    "המשתמש מזין רשימת חומרים שיש לו בבית, "
    "והמערכת מחפשת מתכונים מתאימים בבסיס הנתונים. "
    "אם לא נמצא מתכון — AI יוצר מתכון חדש ומותאם אישית.",
    bold_prefix="חיפוש חכם:"
)
add_bullet(doc,
    "AI יכול לשפר את ניסוח רשימת המרכיבים ולמספר את שלבי ההכנה באופן ברור ומסודר. "
    "לחיצה אחת על כפתור 'שפר' — AI מחזיר גרסה מורחבת ומובנית יותר.",
    bold_prefix="שיפור תוכן:"
)
add_bullet(doc,
    "בדף המתכון, המשתמש יכול להמיר כמויות בין גרמים לכוסות — "
    "פעולה שדורשת הבנה הקשרית ולא רק חישוב מספרי.",
    bold_prefix="המרת יחידות:"
)
add_bullet(doc,
    "כפל כמויות (×1 עד ×10) עם AI מבטיח התאמה נכונה גם של הוראות ההכנה ולא רק המרכיבים.",
    bold_prefix="כפל מתכון:"
)

add_paragraph(doc,
    "כל יכולות ה-AI מופנות דרך AIController בשרת, "
    "המשתמש ב-Groq API (מודל llama-3.3-70b-versatile). "
    "המפתח אינו חשוף בצד הלקוח, "
    "וכל בקשות ה-AI עוברות ולידציה ועיבוד בשרת לפני שהתוצאה מוחזרת."
)

# ─────────────────────────────────────────────
# 12.7 מסקנות
# ─────────────────────────────────────────────
add_subheading(doc, "12.7 מסקנות")

add_paragraph(doc,
    "פרויקט RecipeBox הדגים כיצד ניתן לבנות מערכת רשת מלאה ומאובטחת "
    "תוך שמירה על ארכיטקטורה נקייה ומודולרית. "
    "מספר עקרונות מנחים ליוו את כל תהליך הפיתוח:"
)

add_numbered(doc,
    "הפרדת אחריות — כל שכבה (לקוח, שרת, נתונים) מטפלת בדיוק במה שמוטל עליה, "
    "ואינה חודרת לתחום האחרת."
)
add_numbered(doc,
    "אבטחה בעומק — כל הגנה מיושמת בשכבות מרובות, "
    "כך שכשל חלקי בשכבה אחת אינו חושף את המערכת כולה."
)
add_numbered(doc,
    "ניסיון משתמש (UX) — אפקטים כמו עדכון מועדפים מיידי, "
    "אנימציית הצלחה עם קונפטי ביצירת מתכון, ו-LoginModal שנפתח ללא עזיבת הדף — "
    "כולם נבחרו לשיפור חוויית השימוש ולא מטעמים טכניים בלבד."
)
add_numbered(doc,
    "שירותים חיצוניים כשכבה — Cloudinary ו-Groq AI אינם חלק מהקוד הפנימי, "
    "אלא שירותים הניתנים להחלפה מבלי לשנות את ליבת המערכת."
)
add_numbered(doc,
    "מינימליזם מוכוון תכלית — אין Redux, אין Microservices, אין Event Sourcing. "
    "הכלים שנבחרו תואמים לגודל ולמורכבות הפרויקט בפועל."
)

# ─────────────────────────────────────────────
# 12.8 המלצות להמשך פיתוח
# ─────────────────────────────────────────────
add_subheading(doc, "12.8 המלצות להמשך פיתוח")

add_paragraph(doc,
    "פרויקט RecipeBox מספק תשתית יציבה להרחבות עתידיות. "
    "להלן אפשרויות שיפור שניתן ליישם בהמשך, "
    "מבוססות על מגבלות שזוהו במהלך הפיתוח:"
)

future = [
    ("Refresh Tokens",
     "כיום, עם תפוגת ה-JWT המשתמש מנותק ונדרש להתחברות מחדש. "
     "הוספת Refresh Token תאפשר חידוש הטוקן שקוף למשתמש."),
    ("תגובות ודירוגים",
     "הוספת מודל Rating ו-Comment לטבלת Recipes תאפשר למשתמשים לדרג ולהגיב על מתכונים, "
     "ותהפוך את המערכת לפלטפורמה חברתית."),
    ("הפחתת עומס על AI",
     "כיום כל קריאה ל-AI יוצרת בקשת HTTP חדשה. "
     "שמירת תשובות AI ב-Cache לפי hash של הקלט "
     "תפחית עלויות ותשפר זמני תגובה."),
    ("חיפוש מלא-טקסטי",
     "כיום החיפוש הוא LIKE פשוט. "
     "שדרוג לחיפוש מלא-טקסטי (PostgreSQL Full-Text Search) "
     "ישפר את דיוק תוצאות החיפוש."),
    ("מנגנון ייצוא",
     "הוספת ייצוא מתכונים לקובץ PDF מוגמר (לא דרך window.print) "
     "תשפר את חוויית שמירת מתכונים."),
    ("בדיקות אוטומטיות",
     "הוספת Unit Tests לשכבת ה-Services ו-Integration Tests לנקודות ה-API "
     "תשפר את אמינות המערכת בפיתוח עתידי."),
]

make_table(doc, ["הרחבה מוצעת", "תיאור"], future)

# ─────────────────────────────────────────────
# 12.9 סיכום כללי
# ─────────────────────────────────────────────
add_subheading(doc, "12.9 סיכום כללי")

add_paragraph(doc,
    "מערכת RecipeBox הוקמה מאפס כפרויקט לימודי בעל ערך מעשי. "
    "היא מדגימה שליטה בכלים ובעקרונות הנדרשים לפיתוח מערכות רשת מודרניות: "
    "React, TypeScript, ASP.NET Core, Entity Framework Core, PostgreSQL, "
    "אימות JWT, הרשאות מבוססות תפקיד, שילוב שירותי ענן ובינה מלאכותית."
)

add_paragraph(doc,
    "הפרויקט עמד ביעדיו המרכזיים: "
    "יצירת ממשק משתמש אינטואיטיבי ועשיר, "
    "שרת מאובטח עם לוגיקה עסקית ברורה, "
    "ושילוב AI שמוסיף ערך אמיתי לחוויית המשתמש. "
    "כל פרטי המימוש תועדו בהרחבה בפרקים הקודמים של ספר הפרויקט, "
    "ומשקפים תהליך פיתוח שיטתי ומאורגן."
)

add_paragraph(doc,
    "RecipeBox הוא יותר מאפליקציית מתכונים — "
    "הוא הוכחת יכולת לבניית מערכת מידע מלאה, "
    "מאובטחת, מדרגית ועשירה בתכונות, "
    "תוך שמירה על ארכיטקטורה נקייה ותיעוד מדויק."
)

# ═══════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════
doc.save(DOC_PATH)
print("Section 12 (Final Summary) added successfully.")
