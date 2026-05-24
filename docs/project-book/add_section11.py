from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")

# ─────────────────────────────────────────────
# Helper functions (same style as previous sections)
# ─────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_subheading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return p

def add_sub3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def screen_block(doc, screen_name, route, component, permission,
                 purpose, inputs, buttons, actions, navigation):
    """Write a full screen description block."""
    add_sub3(doc, screen_name)
    rows = [
        ("נתיב (Route)",      route),
        ("קומפוננטה",          component),
        ("הרשאה נדרשת",       permission),
        ("מטרה",              purpose),
        ("שדות קלט",          inputs),
        ("כפתורים ופעולות",   buttons),
        ("פעולות מרכזיות",    actions),
        ("ניווט יוצא",         navigation),
    ]
    make_table(doc, ["מאפיין", "תיאור"], rows)

# ─────────────────────────────────────────────
# PAGE BREAK + SECTION HEADING
# ─────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "11. תיאור מסכים", level=1)
add_paragraph(doc,
    "פרק זה מתאר את כל מסכי האפליקציה כפי שהם מוגדרים בקבצי React. "
    "כל מסך מתועד עם נתיבו, מטרתו, שדות הקלט, הכפתורים, הפעולות שמתרחשות "
    "והדפים שניתן לנווט אליהם ממנו. "
    "האפליקציה מכילה 13 מסכים: 9 ציבוריים/מוגנים ו-4 מסכי ניהול."
)

# Summary table
add_subheading(doc, "11.1 מפת מסכים")
screens_summary = [
    ("דף הבית",             "/",                    "ציבורי",  "HomePage.tsx"),
    ("דף אלבום",            "/albums/:id",          "ציבורי",  "AlbumPage.tsx"),
    ("דף מתכון",            "/recipes/:id",         "ציבורי",  "RecipePage.tsx"),
    ("עוזר AI",             "/ai",                  "ציבורי",  "AIPage.tsx"),
    ("חלון כניסה/הרשמה",    "(Modal)",              "ציבורי",  "LoginModal.tsx"),
    ("הוספת מתכון",         "/recipes/new",         "משתמש",   "AddRecipePage.tsx"),
    ("עריכת מתכון",         "/recipes/:id/edit",    "משתמש",   "EditRecipePage.tsx"),
    ("מועדפים",             "/favorites",           "משתמש",   "FavoritesPage.tsx"),
    ("ניהול – סטטיסטיקות",  "/admin/stats",         "מנהל",    "AdminStatsPage.tsx"),
    ("ניהול – אלבומים",     "/admin/albums",        "מנהל",    "AdminAlbumsPage.tsx"),
    ("ניהול – מתכונים",     "/admin/recipes",       "מנהל",    "AdminRecipesPage.tsx"),
    ("ניהול – משתמשים",     "/admin/users",         "מנהל",    "AdminUsersPage.tsx"),
    ("Navbar",              "(Layout)",             "תמיד",    "Navbar.tsx"),
]
make_table(doc, ["שם מסך", "נתיב", "הרשאה", "קובץ"], screens_summary)

# ─────────────────────────────────────────────
# 11.2 Navbar
# ─────────────────────────────────────────────
add_subheading(doc, "11.2 Navbar – ניווט עליון")
screen_block(
    doc,
    screen_name  = "Navbar.tsx",
    route        = "כל הדפים (Layout)",
    component    = "Navbar",
    permission   = "נטען תמיד; חלקים מותנים בסטטוס כניסה",
    purpose      = "סרגל ניווט קבוע בראש כל דף. מציג לוגו, קישורים ותפריט משתמש.",
    inputs       = "אין שדות קלט",
    buttons      = "לוגו RecipeBox (→ דף הבית), 'מה יש במקרר?' (→ /ai), "
                   "'★ Favorites' (→ /favorites, למחוברים בלבד), "
                   "תפריט Admin עם 4 פריטים (למנהלים בלבד), "
                   "תפריט משתמש עם 'התנתק' (למחוברים), "
                   "כפתורי Login / Register (לא מחוברים)",
    actions      = "logout: מנקה AuthContext ומנתב ל-/. "
                   "openLogin / openRegister: פותח את LoginModal.",
    navigation   = "/ | /ai | /favorites | /admin/* | LoginModal"
)

# ─────────────────────────────────────────────
# 11.3 LoginModal
# ─────────────────────────────────────────────
add_subheading(doc, "11.3 LoginModal – חלון כניסה והרשמה")
screen_block(
    doc,
    screen_name  = "LoginModal.tsx",
    route        = "(Modal – מופיע מעל כל מסך)",
    component    = "LoginModal",
    permission   = "ציבורי – נפתח מכל מקום באפליקציה",
    purpose      = "חלון צף המאפשר כניסה והרשמה מבלי לעזוב את הדף הנוכחי. "
                   "שני מצבים: 'התחברות' ו-'הרשמה', מוחלפים בטאבים בראש החלון.",
    inputs       = "התחברות: אימייל, סיסמה (עם כפתור הצג/הסתר). "
                   "הרשמה: שם משתמש, אימייל, סיסמה, אימות סיסמה.",
    buttons      = "טאב 'התחברות' / 'הרשמה' (מחליף מצב), "
                   "'התחבר' (submit כניסה), "
                   "'הירשם' (submit הרשמה), "
                   "קישור 'הירשם כאן' / 'התחבר כאן' (מחליף טאב), "
                   "לחיצה על הרקע (backdrop) סוגרת את החלון",
    actions      = "login: קורא POST /api/auth/login, שומר JWT ב-localStorage, סוגר מודאל. "
                   "register: קורא POST /api/auth/register עם שם משתמש, אימייל וסיסמה, סוגר מודאל. "
                   "הצגת הודעה 'פג תוקף החיבור' אם sessionExpired=true.",
    navigation   = "נסגר בהצלחה → נשאר באותו הדף (מצב מחובר)"
)

# ─────────────────────────────────────────────
# 11.4 דף הבית
# ─────────────────────────────────────────────
add_subheading(doc, "11.4 דף הבית – HomePage")
screen_block(
    doc,
    screen_name  = "HomePage.tsx",
    route        = "/",
    component    = "HomePage",
    permission   = "ציבורי",
    purpose      = "דף הנחיתה הראשי של האפליקציה. "
                   "מציג Hero banner, רשת אלבומים עם מיון, ורשימת 4 מתכונים שנוספו לאחרונה.",
    inputs       = "תפריט מיון אלבומים: ברירת מחדל / שם א–ת / הכי מתכונים",
    buttons      = "'הרשמי עכשיו' (Hero – לא מחוברים בלבד → פותח LoginModal), "
                   "'+ הוסף מתכון' (→ /recipes/new, למחוברים בלבד), "
                   "כרטיסיות אלבום (AlbumCard – ניתנות ללחיצה → /albums/:id), "
                   "פריטי 'נוספו לאחרונה' (→ /recipes/:id)",
    actions      = "טעינה: GET /api/albums + GET /api/recipes?sortBy=date. "
                   "מיון אלבומים: מחושב בצד לקוח (useMemo) ללא קריאת API.",
    navigation   = "/albums/:id | /recipes/:id | /recipes/new | LoginModal"
)

# ─────────────────────────────────────────────
# 11.5 דף אלבום
# ─────────────────────────────────────────────
add_subheading(doc, "11.5 דף אלבום – AlbumPage")
screen_block(
    doc,
    screen_name  = "AlbumPage.tsx",
    route        = "/albums/:id",
    component    = "AlbumPage",
    permission   = "ציבורי",
    purpose      = "מציג את כל המתכונים של אלבום ספציפי. "
                   "כולל כותרת אלבום, תיאור, מספר מתכונים, שורת חיפוש ומיון, "
                   "ורשת כרטיסיות מתכון.",
    inputs       = "SearchBar: חיפוש חופשי לפי שם/תיאור. "
                   "SortSelect: ברירת מחדל (סדר שרת) / תאריך / שם.",
    buttons      = "'חזרה ← כל האלבומים' (→ /), "
                   "'+ הוסף מתכון' (→ /recipes/new?albumId=:id, למחוברים בלבד), "
                   "כרטיסיות RecipeCard (→ /recipes/:id), "
                   "כפתור לב (FavoriteButton) בכל כרטיסייה",
    actions      = "טעינה: GET /api/albums/:id + GET /api/recipes/album/:id. "
                   "חיפוש ומיון: מחושבים בצד לקוח (useMemo). "
                   "שינוי מועדף: מעדכן state מקומי ללא טעינה מחדש.",
    navigation   = "/ | /recipes/:id | /recipes/new?albumId=:id"
)

# ─────────────────────────────────────────────
# 11.6 דף מתכון
# ─────────────────────────────────────────────
add_subheading(doc, "11.6 דף מתכון – RecipePage")
screen_block(
    doc,
    screen_name  = "RecipePage.tsx",
    route        = "/recipes/:id",
    component    = "RecipePage",
    permission   = "ציבורי (פעולות מסוימות דורשות כניסה)",
    purpose      = "מציג את פרטי המתכון המלאים: שם, תיאור, תמונה, סוג (טקסט/קישור), "
                   "מרכיבים, הוראות הכנה, ופס כלי AI לעיבוד חכם של המתכון.",
    inputs       = "פס AI (מתכוני טקסט בלבד): "
                   "כפתורי 'המר לכוסות' / 'המר לגרמים' (Groq AI), "
                   "כפתורי +/− ו'הכפל' לכפל כמויות (עם בחירת מכפיל ×1–×10), "
                   "כפתור '↩ איפוס' להחזרת ערכים מקוריים.",
    buttons      = "כפתור לב FavoriteButton (הוסף/הסר מועדפים), "
                   "'שתף' (העתק URL ל-clipboard; מחליף ל-'הועתק!'), "
                   "'הורד PDF' (window.print(); דורש כניסה), "
                   "'ערוך' (→ /recipes/:id/edit, לבעלים/מנהל בלבד), "
                   "'מחק' (פותח ConfirmDialog, לבעלים/מנהל בלבד), "
                   "'← חזרה לאלבום' (→ /albums/:albumId), "
                   "Badge אלבום (→ /albums/:albumId), "
                   "כפתור 'פתח מתכון ←' (לפתיחת קישור חיצוני, סוג קישור בלבד)",
    actions      = "טעינה: GET /api/recipes/:id + GET /api/albums/:albumId. "
                   "המרת יחידות / כפל: POST /api/ai/chat דרך geminiService. "
                   "מחיקה: DELETE /api/recipes/:id → navigate לאלבום.",
    navigation   = "/albums/:albumId | /recipes/:id/edit | LoginModal (PDF ללא כניסה)"
)

# ─────────────────────────────────────────────
# 11.7 עוזר AI
# ─────────────────────────────────────────────
add_subheading(doc, "11.7 עוזר AI – AIPage")
screen_block(
    doc,
    screen_name  = "AIPage.tsx",
    route        = "/ai",
    component    = "AIPage",
    permission   = "ציבורי (שמירת מתכון דורשת כניסה)",
    purpose      = "ממשק לחיפוש מתכונים חכם. המשתמש מזין חומרים שיש לו "
                   "ומקבל מתכונים קיימים מהמאגר שמתאימים, או מתכון חדש שנוצר ע\"י AI.",
    inputs       = "Textarea: חומרים שיש למשתמש (או שם מתכון).",
    buttons      = "'חפש מתכון' (submit, מנוטרל אם שדה ריק), "
                   "'צפה במתכון' (→ /recipes/:id, עבור מתכון שנמצא), "
                   "'+ הוסף לאלבום' (פותח מודאל בחירת אלבום, למחוברים), "
                   "'התחבר לשמירה' (פותח LoginModal, לאורחים), "
                   "'✅ נשמר — צפה' (→ /recipes/:savedId, לאחר שמירה). "
                   "במודאל בחירת אלבום: NativeSelect לבחירת אלבום, 'שמור מתכון', 'ביטול'.",
    actions      = "חיפוש: GET /api/recipes (כל המתכונים) → findRecipesWithAI() → POST /api/ai/chat. "
                   "שמירת מתכון: POST /api/recipes עם recipeType='Text'.",
    navigation   = "/recipes/:id (לאחר 'צפה') | /recipes/:savedId (לאחר שמירה) | LoginModal"
)

# ─────────────────────────────────────────────
# 11.8 הוספת מתכון
# ─────────────────────────────────────────────
add_subheading(doc, "11.8 הוספת מתכון – AddRecipePage")
screen_block(
    doc,
    screen_name  = "AddRecipePage.tsx + RecipeForm.tsx",
    route        = "/recipes/new",
    component    = "AddRecipePage (עוטף את RecipeForm)",
    permission   = "משתמש מחובר בלבד (ProtectedRoute)",
    purpose      = "טופס יצירת מתכון חדש. תומך בשני סוגי מתכון: "
                   "טקסט מלא (מרכיבים + הוראות) וקישור חיצוני. "
                   "כולל עוזרי AI ואפשרות הקלטת קול.",
    inputs       = "שם המתכון (חובה), תיאור (אופציונלי), "
                   "בחירת אלבום (NativeSelect, חובה), "
                   "סוג מתכון (כפתורי Toggle: 'טקסט מלא' / 'קישור'), "
                   "אם קישור: שדה URL (חובה). "
                   "אם טקסט: Textarea מרכיבים (חובה), Textarea הוראות הכנה (חובה). "
                   "תמונה: Toggle 'העלאת קובץ' / 'כתובת URL' (אופציונלי).",
    buttons      = "כפתורי סוג מתכון (Toggle), "
                   "כפתורי תמונה Toggle (העלאה/URL), "
                   "כפתור מיקרופון (הקלטת קול ל-Web Speech API, עברית), "
                   "כפתור 'עצור הקלטה', "
                   "כפתור AI 'שפר' (מרכיבים/הוראות → POST /api/ai/chat), "
                   "כפתור '↩ ביטול' (ביטול שיפור AI), "
                   "כפתור AI 'הצע שם ותיאור לפי המרכיבים', "
                   "'צור מתכון' (submit)",
    actions      = "אם יש קובץ תמונה: POST /api/images/upload → Cloudinary. "
                   "שמירה: POST /api/recipes עם JWT. "
                   "הצלחה: הצגת מסך אנימציה עם קונפטי (canvas-confetti), "
                   "ניווט אוטומטי לאלבום לאחר 2.5 שניות.",
    navigation   = "/albums/:albumId (לאחר יצירה)"
)

# ─────────────────────────────────────────────
# 11.9 עריכת מתכון
# ─────────────────────────────────────────────
add_subheading(doc, "11.9 עריכת מתכון – EditRecipePage")
screen_block(
    doc,
    screen_name  = "EditRecipePage.tsx + RecipeForm.tsx",
    route        = "/recipes/:id/edit",
    component    = "EditRecipePage (עוטף את RecipeForm)",
    permission   = "בעלים של המתכון או מנהל (ProtectedRoute + בדיקה בקוד)",
    purpose      = "טופס עריכת מתכון קיים. זהה ל-AddRecipePage בממשק, "
                   "אך שדות מאוכלסים מראש עם ערכי המתכון הנוכחי.",
    inputs       = "זהים לטופס הוספת מתכון, עם ערכים מאוכלסים מראש.",
    buttons      = "זהים לטופס הוספת מתכון. כפתור submit מוצג כ-'שמור שינויים'.",
    actions      = "טעינה: GET /api/recipes/:id. "
                   "אם המשתמש אינו בעלים ואינו מנהל → navigate לדף המתכון. "
                   "שמירה: PUT /api/recipes/:id עם JWT.",
    navigation   = "/recipes/:id (לאחר שמירה)"
)

# ─────────────────────────────────────────────
# 11.10 מועדפים
# ─────────────────────────────────────────────
add_subheading(doc, "11.10 דף מועדפים – FavoritesPage")
screen_block(
    doc,
    screen_name  = "FavoritesPage.tsx",
    route        = "/favorites",
    component    = "FavoritesPage",
    permission   = "משתמש מחובר בלבד (ProtectedRoute)",
    purpose      = "מציג את כל המתכונים שהמשתמש המחובר הוסיף למועדפים, "
                   "בפריסת רשת של כרטיסיות RecipeCard.",
    inputs       = "אין שדות קלט",
    buttons      = "כפתור לב (FavoriteButton) בכל כרטיסייה – לחיצה מסירה את המתכון מהרשימה, "
                   "כרטיסיית RecipeCard (→ /recipes/:id)",
    actions      = "טעינה: GET /api/users/:userId/favorites. "
                   "הסרה: DELETE → מסנן את המתכון מ-state מקומי ללא טעינה מחדש.",
    navigation   = "/recipes/:id"
)

# ─────────────────────────────────────────────
# 11.11 Admin: סטטיסטיקות
# ─────────────────────────────────────────────
add_subheading(doc, "11.11 ניהול – סטטיסטיקות (AdminStatsPage)")
screen_block(
    doc,
    screen_name  = "AdminStatsPage.tsx",
    route        = "/admin/stats",
    component    = "AdminStatsPage",
    permission   = "מנהל בלבד (ProtectedRoute adminOnly)",
    purpose      = "לוח בקרה סטטיסטי עבור המנהל. "
                   "מציג 4 כרטיסי מידע: סה\"כ משתמשים, מתכונים, אלבומים, משתמשים חסומים. "
                   "מציג גם רשימת מתכונים שנוצרו לאחרונה.",
    inputs       = "אין שדות קלט",
    buttons      = "קישורי שמות מתכונים (→ /recipes/:id)",
    actions      = "טעינה: GET /api/users/stats.",
    navigation   = "/recipes/:id (מרשימת המתכונים האחרונים)"
)

# ─────────────────────────────────────────────
# 11.12 Admin: ניהול אלבומים
# ─────────────────────────────────────────────
add_subheading(doc, "11.12 ניהול – אלבומים (AdminAlbumsPage)")
screen_block(
    doc,
    screen_name  = "AdminAlbumsPage.tsx",
    route        = "/admin/albums",
    component    = "AdminAlbumsPage",
    permission   = "מנהל בלבד (ProtectedRoute adminOnly)",
    purpose      = "ניהול מלא של אלבומים: צפייה, יצירה, עריכה ומחיקה. "
                   "כולל טבלת אלבומים וטופס inline לעריכה/יצירה.",
    inputs       = "טופס אלבום (מופיע בלחיצה): שם (חובה), תיאור, תמונת רקע. "
                   "תמונה: שדה URL להדבקה ישירה, או לחצן 'העלה תמונה' לקובץ מהמחשב.",
    buttons      = "'+ אלבום חדש' (פותח טופס ריק), "
                   "כפתור עריכה (עט) בכל שורה (פותח טופס עם ערכים קיימים), "
                   "כפתור מחיקה (פח) בכל שורה (פותח ConfirmDialog), "
                   "'שמור' (submit הטופס), "
                   "'ביטול' (סגירת הטופס), "
                   "כפתור 'העלה / החלף תמונה', 'הסר' תמונה",
    actions      = "יצירה: POST /api/albums. עריכה: PUT /api/albums/:id. "
                   "מחיקה: DELETE /api/albums/:id (מחיקה מדורגת של מתכונים). "
                   "העלאת תמונה: POST /api/images/upload → Cloudinary.",
    navigation   = "אין ניווט יוצא (הכל inline)"
)

# ─────────────────────────────────────────────
# 11.13 Admin: ניהול מתכונים
# ─────────────────────────────────────────────
add_subheading(doc, "11.13 ניהול – מתכונים (AdminRecipesPage)")
screen_block(
    doc,
    screen_name  = "AdminRecipesPage.tsx",
    route        = "/admin/recipes",
    component    = "AdminRecipesPage",
    permission   = "מנהל בלבד (ProtectedRoute adminOnly)",
    purpose      = "מציג טבלת כל מתכוני המערכת עם חיפוש, מיון ואפשרות מחיקה מכל מתכון. "
                   "עריכה מנווטת לדף EditRecipePage הרגיל.",
    inputs       = "SearchBar: חיפוש לפי שם. SortSelect: מיון לפי תאריך / שם.",
    buttons      = "כפתור עריכה (עט) בכל שורה (→ /recipes/:id/edit), "
                   "כפתור מחיקה (פח) בכל שורה (פותח ConfirmDialog), "
                   "קישור שם המתכון (→ /recipes/:id)",
    actions      = "טעינה: GET /api/recipes?search=...&sortBy=... (מופעלת מחדש בכל שינוי). "
                   "מחיקה: DELETE /api/recipes/:id.",
    navigation   = "/recipes/:id | /recipes/:id/edit"
)

# ─────────────────────────────────────────────
# 11.14 Admin: ניהול משתמשים
# ─────────────────────────────────────────────
add_subheading(doc, "11.14 ניהול – משתמשים (AdminUsersPage)")
screen_block(
    doc,
    screen_name  = "AdminUsersPage.tsx",
    route        = "/admin/users",
    component    = "AdminUsersPage",
    permission   = "מנהל בלבד (ProtectedRoute adminOnly)",
    purpose      = "מציג טבלת כל משתמשי המערכת עם סטטוס וכפתור חסימה/ביטול. "
                   "כולל: אימייל, תפקיד (Badge כחול/סגול), מספר מתכונים, סטטוס (ירוק/אדום).",
    inputs       = "אין שדות קלט",
    buttons      = "'חסום' (PATCH /api/users/:id/block, לא מנהלים בלבד), "
                   "'בטל חסימה' (PATCH /api/users/:id/unblock)",
    actions      = "טעינה: GET /api/users. "
                   "חסימה/ביטול: PATCH → מעדכן state מקומי ללא טעינה מחדש.",
    navigation   = "אין ניווט יוצא"
)

# ─────────────────────────────────────────────
# 11.15 RecipeForm – כרטיס AI
# ─────────────────────────────────────────────
add_subheading(doc, "11.15 RecipeForm – יכולות AI ותמיכה בקול")
add_paragraph(doc,
    "RecipeForm הוא קומפוננט משותף בין AddRecipePage ו-EditRecipePage. "
    "מעבר לשדות הטופס הרגילים, הוא כולל שלוש יכולות מתקדמות:"
)
add_bullet(doc,
    "הקלטת קול (Web Speech API): כפתור מיקרופון ליד שדות 'מרכיבים' ו'הוראות הכנה'. "
    "פועל בעברית (he-IL), מוסיף טקסט בזמן אמת. "
    "סף ברירת מחדל לדפדפנים התומכים (Chrome בעיקר).",
    bold_prefix="הקלטת קול:"
)
add_bullet(doc,
    "שיפור טקסט AI: כפתור '✨ שפר' בכל שדה טקסט. שולח את התוכן ל-POST /api/ai/chat (Groq). "
    "מחזיר גרסה משופרת; כפתור '↩ ביטול' משחזר את הטקסט המקורי.",
    bold_prefix="שיפור AI:"
)
add_bullet(doc,
    "הצעת שם ותיאור: כפתור '✨ הצע שם ותיאור לפי המרכיבים' (מופיע רק בטקסט). "
    "שולח את המרכיבים ל-AI ומאכלס אוטומטית שדות שם ותיאור.",
    bold_prefix="הצעת שם:"
)

# ─────────────────────────────────────────────
# 11.16 סיכום
# ─────────────────────────────────────────────
add_subheading(doc, "11.16 סיכום")
add_paragraph(doc,
    "האפליקציה כוללת 13 מסכים המחולקים לשלוש רמות גישה: "
    "4 מסכים ציבוריים (דף הבית, אלבום, מתכון, AI), "
    "3 מסכים למשתמשים מחוברים (הוספה, עריכה, מועדפים), "
    "ו-4 מסכי ניהול בלעדיים למנהלים. "
    "LoginModal הוא חלון צף גלובלי המאפשר כניסה והרשמה מכל עמוד ב-SPA "
    "מבלי לאבד את ההקשר הנוכחי. "
    "ה-Navbar נטען בכל דף כחלק מה-Layout ומציג תפריטים מותאמים לסטטוס הכניסה."
)

doc.save("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")
print("Section 11 added successfully.")
