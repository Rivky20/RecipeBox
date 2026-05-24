from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_subheading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return p

def add_sub3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ===========================
# PAGE BREAK
# ===========================
doc.add_page_break()

# ===========================
# SECTION 8 HEADING
# ===========================
add_heading(doc, "8. פונקציות ולוגיקה מרכזית", level=1)

add_paragraph(doc,
    "פרק זה מתעד את הלוגיקה העסקית המרכזית של מערכת RecipeBox, "
    "כפי שהיא מיושמת בקוד הממשי של השרת (C# / ASP.NET Core) ושל הלקוח (React / TypeScript). "
    "כל הפרטים מבוססים אך ורק על קריאת הקוד בפועל."
)

# ===========================
# 8.1 מבנה ה-API
# ===========================
add_subheading(doc, "8.1 מבנה ה-API")

add_paragraph(doc,
    "ה-API בנוי כ-RESTful API תחת בסיס הנתיב /api. "
    "כל בקר (Controller) מקבל את התלויות שלו דרך Dependency Injection. "
    "ה-Enum RecipeType מועבר בין הלקוח לשרת כמחרוזת (\"Text\" / \"Link\") "
    "בזכות JsonStringEnumConverter שמוגדר ב-Program.cs."
)

api_routes = [
    ("POST",   "/api/auth/register",                "ציבורי",       "רישום משתמש חדש"),
    ("POST",   "/api/auth/login",                   "ציבורי",       "כניסה עם אימייל וסיסמה"),
    ("GET",    "/api/recipes",                      "ציבורי",       "שליפת כל המתכונים (עם פילטור ומיון)"),
    ("GET",    "/api/recipes/{id}",                 "ציבורי",       "שליפת מתכון לפי מזהה"),
    ("GET",    "/api/recipes/album/{albumId}",      "ציבורי",       "שליפת מתכונים לפי אלבום"),
    ("POST",   "/api/recipes",                      "מחובר",        "יצירת מתכון חדש"),
    ("PUT",    "/api/recipes/{id}",                 "מחובר",        "עדכון מתכון (בעלים בלבד, או Admin)"),
    ("DELETE", "/api/recipes/{id}",                 "מחובר",        "מחיקת מתכון (בעלים בלבד, או Admin)"),
    ("GET",    "/api/albums",                       "ציבורי",       "שליפת כל האלבומים"),
    ("GET",    "/api/albums/{id}",                  "ציבורי",       "שליפת אלבום לפי מזהה"),
    ("POST",   "/api/albums",                       "Admin בלבד",   "יצירת אלבום חדש"),
    ("PUT",    "/api/albums/{id}",                  "Admin בלבד",   "עדכון אלבום"),
    ("DELETE", "/api/albums/{id}",                  "Admin בלבד",   "מחיקת אלבום"),
    ("GET",    "/api/users/{userId}/favorites",     "מחובר (עצמי)", "שליפת מועדפי משתמש"),
    ("POST",   "/api/users/{userId}/favorites",     "מחובר (עצמי)", "הוספת מתכון למועדפים"),
    ("DELETE", "/api/users/{userId}/favorites/{id}","מחובר (עצמי)", "הסרת מתכון מהמועדפים"),
    ("GET",    "/api/users",                        "Admin בלבד",   "שליפת כל המשתמשים"),
    ("GET",    "/api/users/{id}",                   "Admin בלבד",   "שליפת משתמש לפי מזהה"),
    ("PATCH",  "/api/users/{id}/block",             "Admin בלבד",   "חסימת משתמש"),
    ("PATCH",  "/api/users/{id}/unblock",           "Admin בלבד",   "ביטול חסימת משתמש"),
    ("GET",    "/api/users/stats",                  "Admin בלבד",   "שליפת סטטיסטיקות מערכת"),
    ("POST",   "/api/ai/chat",                      "ציבורי",       "פרוקסי לשירות Groq AI"),
    ("POST",   "/api/images/upload",                "מחובר",        "העלאת תמונה ל-Cloudinary"),
]
make_table(doc,
    ["Method", "נתיב", "הרשאה נדרשת", "תיאור"],
    api_routes
)

# ===========================
# 8.2 לוגיקת אימות (Authentication)
# ===========================
add_subheading(doc, "8.2 לוגיקת אימות (Authentication)")

add_sub3(doc, "8.2.1 רישום משתמש – RegisterAsync")

add_paragraph(doc,
    "המתודה RegisterAsync ב-AuthService מבצעת את הפעולות הבאות לפי הסדר:"
)
add_bullet(doc, "בדיקת ייחודיות האימייל: שאילתה ל-DB לפי Email.ToLower() – אם קיים, נזרקת InvalidOperationException.")
add_bullet(doc, "יצירת אובייקט User עם Role='User' ו-PasswordHash שנוצר על-ידי BCrypt.Net.BCrypt.HashPassword().")
add_bullet(doc, "שמירה ב-DB באמצעות SaveChangesAsync.")
add_bullet(doc, "יצירה והחזרת AuthResponseDto הכולל: Token (JWT), UserId, Email, Role, UserName.")

add_sub3(doc, "8.2.2 כניסה – LoginAsync")

add_paragraph(doc,
    "המתודה LoginAsync מבצעת:"
)
add_bullet(doc, "שליפת משתמש לפי Email.ToLower() מה-DB.")
add_bullet(doc, "אם המשתמש לא נמצא, או שסיסמת הבדיקה נכשלת (BCrypt.Verify) – נזרקת UnauthorizedAccessException.")
add_bullet(doc, "אם השדה IsBlocked=true – נזרקת UnauthorizedAccessException עם הודעה 'Your account has been blocked.'")
add_bullet(doc, "אם הכניסה תקינה – מוחזר AuthResponseDto עם JWT חדש.")

add_sub3(doc, "8.2.3 יצירת טוקן JWT – GenerateToken")

add_paragraph(doc,
    "המתודה הפרטית GenerateToken יוצרת טוקן JWT ולא מוחזרת כ-API ישיר:"
)
add_bullet(doc, "מפתח ההצפנה נקרא מתוך appsettings.json תחת Jwt:Key.")
add_bullet(doc, "אלגוריתם החתימה: HmacSha256.")
add_bullet(doc, "Claims הכלולים: ClaimTypes.NameIdentifier (UserId), ClaimTypes.Email, ClaimTypes.Role.")
add_bullet(doc, "תוקף הטוקן: 7 ימים מרגע היצירה (DateTime.UtcNow.AddDays(7)).")
add_bullet(doc, "Issuer ו-Audience אינם מאומתים (ValidateIssuer=false, ValidateAudience=false).")
add_bullet(doc, "ClockSkew מוגדר לאפס – הטוקן פג בדיוק בזמן שצוין.")

# ===========================
# 8.3 לוגיקת הרשאה (Authorization)
# ===========================
add_subheading(doc, "8.3 לוגיקת הרשאה (Authorization)")

add_sub3(doc, "8.3.1 רמות הרשאה בשרת")

add_paragraph(doc,
    "המערכת מיישמת שלוש רמות הרשאה בשרת:"
)
add_bullet(doc, "ציבורי – אנדפוינטים ללא [Authorize]; נגישים לכל משתמש, כולל לא מחובר.",
           bold_prefix="ציבורי:")
add_bullet(doc,
    "מחובר – מסומן ב-[Authorize]; דורש JWT תקין ב-Authorization header. "
    "הבקרים RecipesController ו-FavoritesController משתמשים בזה.",
    bold_prefix="מחובר:")
add_bullet(doc,
    "Admin בלבד – מסומן ב-[Authorize(Roles = \"Admin\")]; דורש שה-Claim Role יהיה 'Admin'. "
    "כל UsersController וכן POST/PUT/DELETE ב-AlbumsController מוגנים כך.",
    bold_prefix="Admin בלבד:")

add_sub3(doc, "8.3.2 הרשאת בעלות על משאב (Ownership Check)")

add_paragraph(doc,
    "בעת עדכון או מחיקת מתכון, RecipeService בודק בעלות:"
)
add_bullet(doc, "אם המשתמש אינו Admin ו-recipe.UserId != userId – נזרקת UnauthorizedAccessException.")
add_bullet(doc, "מנהל מערכת (isAdmin=true) יכול לערוך ולמחוק כל מתכון.")
add_bullet(doc, "הפרמטר isAdmin מועבר מהבקר: var isAdmin = User.IsInRole(\"Admin\").")

add_sub3(doc, "8.3.3 הרשאת גישה למועדפים (CanAccessUser)")

add_paragraph(doc,
    "FavoritesController מגן על כל endpoint עם בדיקה פרטית:"
)
add_bullet(doc,
    "CanAccessUser(userId) מחלצת את ה-ClaimTypes.NameIdentifier מהטוקן "
    "ומשווה אותה ל-userId שב-URL. "
    "גישה מותרת רק אם currentId == userId, או אם המשתמש הוא Admin."
)

add_sub3(doc, "8.3.4 הגנת מסלולים בצד הלקוח – ProtectedRoute")

add_paragraph(doc,
    "קומפוננט ProtectedRoute בצד הלקוח מגן על מסלולים הדורשים אימות:"
)
add_bullet(doc, "אם isAuthenticated=false – נפתח חלון ModalLogin ומבוצע redirect ל-/.")
add_bullet(doc, "אם adminOnly=true ו-isAdmin=false – redirect ל-/ ללא הודעה.")
add_bullet(doc, "המסלולים /recipes/new, /recipes/:id/edit ו-/favorites מוגנים ב-ProtectedRoute.")
add_bullet(doc, "המסלולים /admin/* מוגנים ב-<ProtectedRoute adminOnly>.")

# ===========================
# 8.4 פעולות CRUD
# ===========================
add_subheading(doc, "8.4 פעולות CRUD")

add_sub3(doc, "8.4.1 מתכונים (RecipeService)")

add_paragraph(doc, "יצירת מתכון – CreateAsync:", bold_prefix="")
add_bullet(doc, "קריאה ל-ValidateRecipeType: מתכון מסוג Link חייב Link; מסוג Text חייב Ingredients או Instructions.")
add_bullet(doc, "בדיקת קיום האלבום (AlbumId) ב-DB לפני יצירה.")
add_bullet(doc, "שדות שאינם רלוונטיים לסוג המתכון מוגדרים null (Link=null עבור Text, ולהפך).")
add_bullet(doc, "לאחר SaveChangesAsync, נטענים Navigation Properties (Album, User) לפני המרה ל-DTO.")

add_paragraph(doc, "עדכון מתכון – UpdateAsync:", bold_prefix="")
add_bullet(doc, "שליפת המתכון עם Include של Album, User, FavoritedBy.")
add_bullet(doc, "בדיקת בעלות; בדיקת קיום אלבום חדש; ולידציה של סוג מתכון.")
add_bullet(doc, "עדכון כל השדות ישירות על האובייקט ואחר-כך SaveChangesAsync.")

add_paragraph(doc, "מחיקת מתכון – DeleteAsync:", bold_prefix="")
add_bullet(doc, "שליפה ב-FindAsync (ללא Include – לא נדרש לוולידציה).")
add_bullet(doc, "בדיקת בעלות; קריאה ל-_db.Recipes.Remove; SaveChangesAsync.")
add_bullet(doc, "מחיקת UserFavorites הקשורות מתבצעת אוטומטית (Cascade).")

add_paragraph(doc, "מיון מתכונים – ApplySorting:", bold_prefix="")
add_bullet(doc,
    "אם המשתמש מחובר: מתכוני המשתמש עצמו קודם, "
    "לאחר-מכן מועדפים, ואחרון לפי תאריך יצירה יורד."
)
add_bullet(doc,
    "אם המשתמש אינו מחובר: מיון לפי שם (SortBy=name) "
    "או לפי תאריך יצירה יורד (ברירת מחדל)."
)

add_sub3(doc, "8.4.2 אלבומים (AlbumService)")

crud_album = [
    ("GetAllAsync",   "שליפת כל האלבומים עם Include(Recipes). מיון לפי Name. מוחזר RecipeCount מחושב."),
    ("GetByIdAsync",  "שליפת אלבום יחיד עם Include(Recipes). מוחזר null אם לא נמצא."),
    ("CreateAsync",   "יצירת Album חדש עם Name, Description, ImagePath. SaveChangesAsync ולאחר-מכן מוחזר DTO עם RecipeCount=0."),
    ("UpdateAsync",   "FindAsync; עדכון Name, Description, ImagePath; SaveChangesAsync. ספירת Recipes מחושבת מחדש."),
    ("DeleteAsync",   "FindAsync; Remove; SaveChangesAsync. מחיקת מתכוני האלבום מתבצעת אוטומטית (Cascade)."),
]
make_table(doc, ["מתודה", "תיאור"], crud_album)

add_sub3(doc, "8.4.3 מועדפים (FavoriteService)")

add_paragraph(doc,
    "ניהול מועדפים מבוצע דרך FavoriteService דרך טבלת UserFavorites:"
)
add_bullet(doc, "GetFavoritesAsync – שליפת כל המתכונים המועדפים של משתמש, עם Include של Album, User, FavoritedBy.")
add_bullet(doc, "AddFavoriteAsync – בדיקת כפילות; אם קיים כבר מוחזר false (Conflict). אחרת נוצרת רשומת UserFavorite ומוחזר true.")
add_bullet(doc, "RemoveFavoriteAsync – שליפה; אם לא נמצא מוחזר false. אחרת Remove ו-SaveChangesAsync.")

# ===========================
# 8.5 ניהול מצב (State Management)
# ===========================
add_subheading(doc, "8.5 ניהול מצב (State Management)")

add_paragraph(doc,
    "המערכת אינה משתמשת ב-Redux. ניהול המצב מבוסס על React Context API בלבד."
)

add_sub3(doc, "8.5.1 AuthContext – מצב האימות")

add_paragraph(doc,
    "AuthContext מנהל את מצב המשתמש המחובר בכל האפליקציה:"
)
add_bullet(doc,
    "המצב ההתחלתי נטען מ-localStorage תחת המפתח 'recipebox_user' "
    "בפונקציית האתחול של useState – כך שהמשתמש נשאר מחובר לאחר רענון הדפדפן."
)
add_bullet(doc,
    "useEffect מסנכרן את state עם localStorage בכל שינוי: "
    "אם user!=null שומר JSON, אם null מסיר את הפריט."
)
add_bullet(doc,
    "useEffect מאזין לאירוע 'session-expired' על window – "
    "כאשר Axios מקבל 401, הוא שולח אירוע זה, "
    "ו-AuthContext מאפס את המשתמש ל-null אוטומטית."
)
add_bullet(doc, "login – קריאה ל-authService.login, שמירת המשתמש ב-state.")
add_bullet(doc, "register – קריאה ל-authService.register, שמירת המשתמש ב-state.")
add_bullet(doc, "logout – מאפס את state ל-null (ה-useEffect מנקה את localStorage).")
add_bullet(doc, "isAdmin – derived value: user?.role === 'Admin'.")

add_sub3(doc, "8.5.2 AuthModalContext – מצב חלון ההתחברות")

add_paragraph(doc,
    "AuthModalContext מנהל את פתיחת/סגירת חלון ה-Modal של ההתחברות והרישום. "
    "חשיפת useAuthModal מאפשרת לכל קומפוננט לפתוח את המודל ללא props drilling."
)

# ===========================
# 8.6 טפסים ואימות קלט (Forms & Validation)
# ===========================
add_subheading(doc, "8.6 טפסים ואימות קלט")

add_sub3(doc, "8.6.1 אימות צד-לקוח – RecipeForm")

add_paragraph(doc,
    "פונקציית validate() ב-RecipeForm מחזירה מילון שגיאות:"
)
add_bullet(doc, "name – שדה חובה; אם ריק: 'שם המתכון נדרש.'")
add_bullet(doc, "albumId – חייב להיות שונה מ-0; אם לא נבחר: 'נא לבחור אלבום.'")
add_bullet(doc, "link – חובה רק אם recipeType='Link'; אם ריק: 'כתובת הקישור נדרשת.'")
add_bullet(doc, "ingredients – חובה רק אם recipeType='Text'; אם ריק: 'נא להזין מרכיבים.'")
add_bullet(doc, "instructions – חובה רק אם recipeType='Text'; אם ריק: 'נא להזין הוראות הכנה.'")
add_paragraph(doc,
    "אם מילון השגיאות אינו ריק, הפורם אינו נשלח ושגיאות מוצגות ליד כל שדה "
    "דרך Field.ErrorText של Chakra UI."
)

add_sub3(doc, "8.6.2 אימות צד-שרת – ValidateRecipeType")

add_paragraph(doc,
    "גם בשרת מתבצעת ולידציה דרך ValidateRecipeType (נקראת ב-CreateAsync ו-UpdateAsync):"
)
add_bullet(doc, "RecipeType.Link ללא Link → InvalidOperationException('Link is required for Link-type recipes.')")
add_bullet(doc, "RecipeType.Text ללא Ingredients וגם ללא Instructions → InvalidOperationException.")
add_bullet(doc, "AlbumId שאינו קיים ב-DB → InvalidOperationException('Album not found.')")

add_sub3(doc, "8.6.3 העלאת תמונה")

add_paragraph(doc,
    "ב-RecipeForm ישנן שתי שיטות להוספת תמונה:"
)
add_bullet(doc,
    "העלאת קובץ (imageMode='upload'): "
    "המשתמש בוחר קובץ, imageService.upload() שולח אותו ל-/api/images/upload, "
    "השרת מעלה ל-Cloudinary ומחזיר URL ציבורי."
)
add_bullet(doc,
    "כתובת URL (imageMode='url'): "
    "המשתמש מדביק URL ישירות; הוא נשמר ב-imagePath ללא העלאה."
)
add_bullet(doc,
    "במתכון מסוג Link – שדה ה-URL לתמונה מוצג תמיד (ללא אפשרות העלאת קובץ), "
    "מכיוון שהמתכון מצביע לאתר חיצוני."
)

# ===========================
# 8.7 טיפול בשגיאות (Error Handling)
# ===========================
add_subheading(doc, "8.7 טיפול בשגיאות")

add_sub3(doc, "8.7.1 שרת – בקרים")

add_paragraph(doc,
    "כל הבקרים משתמשים ב-try/catch לתרגום חריגות לקודי HTTP:"
)

error_map = [
    ("InvalidOperationException", "400 Bad Request",   "קלט לא תקין (חסר שדה חובה, אלבום לא קיים, וכדומה)"),
    ("UnauthorizedAccessException","401 Unauthorized / 403 Forbidden",
     "פרטי כניסה שגויים (401) או גישה לא מורשית למשאב (403/Forbid)"),
    ("null מה-Service",           "404 Not Found",     "משאב לא נמצא (מתכון, אלבום, משתמש)"),
    ("כפילות (false מ-Service)",  "409 Conflict",      "מועדף שכבר קיים; אימייל שכבר רשום"),
]
make_table(doc, ["חריגה / מצב", "קוד HTTP", "מקרה לדוגמה"], error_map)

add_sub3(doc, "8.7.2 לקוח – Axios Interceptors")

add_paragraph(doc,
    "קובץ axios.ts מגדיר שני Interceptors המרכזים את הטיפול בשגיאות:"
)
add_bullet(doc,
    "Request Interceptor: קריאת הטוקן מ-localStorage ב-'recipebox_user' "
    "והוספתו כ-Authorization: Bearer <token> לכל בקשה יוצאת."
)
add_bullet(doc,
    "Response Interceptor: אם תגובת השרת היא 401 – "
    "מסיר את 'recipebox_user' מ-localStorage "
    "ושולח CustomEvent('session-expired') שמאזין אליו AuthContext "
    "ומאפס את המשתמש."
)

add_sub3(doc, "8.7.3 לקוח – טפסים")

add_paragraph(doc,
    "בכל טופס (RecipeForm, LoginModal, RegisterPage) "
    "שגיאות API מוצגות בתוך Alert.Root של Chakra UI "
    "מתחת לכותרת הטופס. "
    "הודעת השגיאה נלקחת מ-err?.response?.data?.message, "
    "ואם אינה קיימת – מ-err?.message, "
    "ואם גם זה אינו קיים – מוצגת הודעת ברירת מחדל."
)

# ===========================
# 8.8 שירות בינה מלאכותית
# ===========================
add_subheading(doc, "8.8 שירות הבינה המלאכותית")

add_paragraph(doc,
    "המערכת כוללת אינטגרציה עם שני שירותי AI שונים, "
    "כאשר כל הגישה לשירותים האלה עוברת דרך פרוקסי בשרת (AIController)."
)

add_sub3(doc, "8.8.1 Groq – דרך AIController")

add_paragraph(doc,
    "AIController ב-/api/ai/chat משמש כפרוקסי לשירות Groq AI (מודל llama-3.3-70b-versatile). "
    "המפתח (Groq:ApiKey) נשמר ב-appsettings ואינו חשוף ללקוח. "
    "הלקוח שולח את body הבקשה כפי שהוא, "
    "והבקר מעביר אותו ישירות ל-Groq ומחזיר את התשובה."
)

ai_funcs = [
    ("findRecipesWithAI",       "מחפש מתכונים קיימים לפי מרכיבים שהוזנו, או יוצר מתכון חדש אם לא נמצאו מתאימים. שולח עד 30 מתכוני Text ל-AI."),
    ("improveText",             "משפר טקסט של מרכיבים (הוספת כמויות, ניסוח מסודר) או הוראות הכנה (מספור שלבים, הבהרה)."),
    ("suggestNameAndDescription","מציע שם ותיאור קצר למתכון בעברית, לפי רשימת המרכיבים שהוזנה."),
    ("convertUnits",            "ממיר כמויות ברשימת מרכיבים בין גרמים לכוסות, ולהפך."),
    ("multiplyRecipe",          "מכפיל את כמויות המרכיבים וההוראות במקדם שנבחר על-ידי המשתמש."),
]
make_table(doc, ["פונקציה (geminiService.ts)", "תיאור"], ai_funcs)

add_paragraph(doc,
    "כל פונקציות ה-AI ב-geminiService.ts משתמשות בפונקציה הפנימית groqChat() "
    "שמבצעת fetch ל-/api/ai/chat ומחזירה את תוכן ה-message מהתגובה. "
    "תגובות ה-AI המכילות JSON עוברות ניקוי של קוד markdown (```json```) לפני parse."
)

# ===========================
# 8.9 סיכום
# ===========================
add_subheading(doc, "8.9 סיכום")

add_paragraph(doc,
    "הלוגיקה המרכזית של RecipeBox מחולקת בצורה ברורה בין שכבות: "
    "הבקרים (Controllers) אחראים לתרגום HTTP ולחילוץ פרטי המשתמש מהטוקן; "
    "שכבת ה-Services מרכזת את כל הלוגיקה העסקית, הוולידציה ושאילתות ה-DB; "
    "ו-AuthContext בצד הלקוח מנהל את מצב האימות באמצעות Context API ו-localStorage. "
    "כל הגנות ההרשאה מיושמות גם בשרת (Authorize, ownership check) "
    "וגם בלקוח (ProtectedRoute), "
    "וכל שגיאות ה-API מטופלות באופן אחיד דרך Axios Interceptors."
)

doc.save("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")
print("Section 8 added successfully.")
