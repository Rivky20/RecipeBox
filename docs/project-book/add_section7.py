from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_subheading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return p

def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    # data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ===========================
# PAGE BREAK
# ===========================
doc.add_page_break()

# ===========================
# SECTION 7 HEADING
# ===========================
add_heading(doc, "7. מבנה נתונים", level=1)

add_paragraph(doc,
    "פרק זה מתאר את מבנה הנתונים של מערכת RecipeBox כפי שהוגדר בקבצי המודלים "
    "ובמיגרציות של Entity Framework Core. מסד הנתונים הוא PostgreSQL, "
    "וכל הטבלאות, השדות, המפתחות והקשרים מוצגים להלן על בסיס הקוד בלבד."
)

# ===========================
# 7.1 סקירת הטבלאות
# ===========================
add_subheading(doc, "7.1 סקירת הטבלאות במסד הנתונים")

add_paragraph(doc,
    "מסד הנתונים מכיל ארבע טבלאות מרכזיות:"
)

overview = [
    ("Users",        "טבלת משתמשי המערכת"),
    ("Albums",       "טבלת האלבומים לאיסוף מתכונים"),
    ("Recipes",      "טבלת המתכונים"),
    ("UserFavorites","טבלת מועדפים – קשר רבים-לרבים בין משתמשים למתכונים"),
]
make_table(doc, ["שם הטבלה", "תיאור"], overview)

# ===========================
# 7.2 טבלת Users
# ===========================
add_subheading(doc, "7.2 טבלת Users – משתמשים")

add_paragraph(doc,
    "מכילה את נתוני כלל המשתמשים הרשומים במערכת. "
    "כל משתמש מזוהה על-ידי GUID ייחודי. "
    "שדה האימייל הוא ייחודי (UNIQUE INDEX) ומשמש כמזהה כניסה. "
    "הסיסמה אינה נשמרת כטקסט חופשי אלא כגיבוב BCrypt."
)

users_fields = [
    ("Id",           "uuid",                      "מפתח ראשי (PK)",          "לא",       "ברירת מחדל: Guid.NewGuid()"),
    ("Email",        "text",                      "אימייל המשתמש",            "לא",       "חייב להיות ייחודי (UNIQUE)"),
    ("UserName",     "text",                      "שם משתמש לתצוגה",          "כן (null)","נוסף במיגרציה AddUserName"),
    ("PasswordHash", "text",                      "גיבוב הסיסמה (BCrypt)",    "לא",       "לא נחשף ב-API"),
    ("Role",         "text",                      "תפקיד: 'User' / 'Admin'", "לא",       "ברירת מחדל: 'User'"),
    ("IsBlocked",    "boolean",                   "האם המשתמש חסום",          "לא",       "ברירת מחדל: false"),
    ("CreatedAt",    "timestamp with time zone",  "תאריך יצירת החשבון",       "לא",       "UTC אוטומטי"),
]
make_table(doc,
    ["שם השדה", "סוג נתון", "תיאור", "ניתן לריק?", "הערות"],
    users_fields
)

add_paragraph(doc, "אינדקסים:")
add_bullet(doc, "IX_Users_Email – UNIQUE INDEX על שדה Email")

# ===========================
# 7.3 טבלת Albums
# ===========================
add_subheading(doc, "7.3 טבלת Albums – אלבומים")

add_paragraph(doc,
    "מכילה את האלבומים שמשמשים לאיסוף וסיווג מתכונים. "
    "ניהול האלבומים מוגבל למנהל המערכת בלבד. "
    "שדה ImagePath נוסף במיגרציה UpdateModel לאחר היצירה הראשונית."
)

albums_fields = [
    ("Id",          "integer (auto-increment)", "מפתח ראשי (PK)",      "לא",       "Identity Column – PostgreSQL"),
    ("Name",        "text",                     "שם האלבום",            "לא",       ""),
    ("Description", "text",                     "תיאור האלבום",         "לא",       ""),
    ("ImagePath",   "text",                     "נתיב תמונת האלבום",    "כן (null)","URL מ-Cloudinary"),
    ("CreatedAt",   "timestamp with time zone", "תאריך יצירה",          "לא",       "UTC אוטומטי"),
]
make_table(doc,
    ["שם השדה", "סוג נתון", "תיאור", "ניתן לריק?", "הערות"],
    albums_fields
)

# ===========================
# 7.4 טבלת Recipes
# ===========================
add_subheading(doc, "7.4 טבלת Recipes – מתכונים")

add_paragraph(doc,
    "טבלה מרכזית המכילה את כל המתכונים במערכת. "
    "כל מתכון שייך לאלבום אחד ולמשתמש אחד. "
    "שדה RecipeType נשמר כטקסט (\"Link\" או \"Text\") בהמרה מ-Enum בקוד. "
    "שדות Link, Ingredients ו-Instructions הם אופציונליים ומשמשים לפי סוג המתכון."
)

recipes_fields = [
    ("Id",           "integer (auto-increment)", "מפתח ראשי (PK)",            "לא",       "Identity Column"),
    ("Name",         "text",                     "שם המתכון",                  "לא",       ""),
    ("Description",  "text",                     "תיאור המתכון",               "לא",       ""),
    ("ImagePath",    "text",                     "נתיב תמונה (Cloudinary URL)","כן (null)",""),
    ("RecipeType",   "text",                     "סוג מתכון: 'Link' / 'Text'", "לא",       "נשמר כטקסט, לא כמספר"),
    ("Link",         "text",                     "URL לאתר מתכון חיצוני",      "כן (null)","בשימוש רק אם RecipeType='Link'"),
    ("Ingredients",  "text",                     "רשימת מצרכים",               "כן (null)","בשימוש רק אם RecipeType='Text'"),
    ("Instructions", "text",                     "הוראות הכנה",                "כן (null)","בשימוש רק אם RecipeType='Text'"),
    ("AlbumId",      "integer",                  "מפתח זר → Albums.Id",        "לא",       "ON DELETE CASCADE"),
    ("UserId",       "uuid",                     "מפתח זר → Users.Id",         "לא",       "ON DELETE CASCADE"),
    ("CreatedAt",    "timestamp with time zone", "תאריך יצירת המתכון",         "לא",       "UTC אוטומטי"),
]
make_table(doc,
    ["שם השדה", "סוג נתון", "תיאור", "ניתן לריק?", "הערות"],
    recipes_fields
)

add_paragraph(doc, "אינדקסים:")
add_bullet(doc, "IX_Recipes_AlbumId – אינדקס על AlbumId לשיפור שאילתות סינון לפי אלבום")
add_bullet(doc, "IX_Recipes_UserId  – אינדקס על UserId לשיפור שאילתות שליפת מתכוני משתמש")

# ===========================
# 7.5 טבלת UserFavorites
# ===========================
add_subheading(doc, "7.5 טבלת UserFavorites – מועדפים")

add_paragraph(doc,
    "טבלת חיבור (Junction Table) המממשת קשר רבים-לרבים בין משתמשים למתכונים. "
    "המפתח הראשי הוא מורכב משני השדות יחד: (UserId, RecipeId). "
    "מחיקת משתמש או מתכון מוחקת אוטומטית את כל רשומות המועדפים הקשורות."
)

favorites_fields = [
    ("UserId",    "uuid",                     "מפתח זר → Users.Id (חלק מ-PK מורכב)",   "לא", "ON DELETE CASCADE"),
    ("RecipeId",  "integer",                  "מפתח זר → Recipes.Id (חלק מ-PK מורכב)", "לא", "ON DELETE CASCADE"),
    ("AddedAt",   "timestamp with time zone", "תאריך הוספה למועדפים",                  "לא", "UTC אוטומטי"),
]
make_table(doc,
    ["שם השדה", "סוג נתון", "תיאור", "ניתן לריק?", "הערות"],
    favorites_fields
)

add_paragraph(doc, "מפתח ראשי מורכב:", bold_prefix="")
add_bullet(doc, "PK_UserFavorites = { UserId, RecipeId } – כל שילוב של משתמש ומתכון ייחודי")

add_paragraph(doc, "אינדקסים:")
add_bullet(doc, "IX_UserFavorites_UserId  – לשיפור שאילתת שליפת מועדפי משתמש")
add_bullet(doc, "IX_UserFavorites_RecipeId – לשיפור שאילתת בדיקת מי סימן מתכון כמועדף")

# ===========================
# 7.6 מפתחות זרים וקשרים
# ===========================
add_subheading(doc, "7.6 מפתחות זרים וקשרים בין הטבלאות")

fk_rows = [
    ("FK_Recipes_Users_UserId",           "Recipes.UserId",        "Users.Id",         "CASCADE", "מחיקת משתמש מוחקת את כל מתכוניו"),
    ("FK_Recipes_Albums_AlbumId",         "Recipes.AlbumId",       "Albums.Id",        "CASCADE", "מחיקת אלבום מוחקת את כל מתכוניו"),
    ("FK_UserFavorites_Users_UserId",     "UserFavorites.UserId",  "Users.Id",         "CASCADE", "מחיקת משתמש מוחקת את כל מועדפיו"),
    ("FK_UserFavorites_Recipes_RecipeId", "UserFavorites.RecipeId","Recipes.Id",       "CASCADE", "מחיקת מתכון מוחקת את כל סימוני המועדפים שלו"),
]
make_table(doc,
    ["שם המפתח הזר", "טבלה.עמודה (מקור)", "טבלה.עמודה (יעד)", "On Delete", "הסבר"],
    fk_rows
)

# ===========================
# 7.7 דיאגרמת קשרים
# ===========================
add_subheading(doc, "7.7 תיאור מילולי של הקשרים")

relationships = [
    ("User → Recipe (1:N)",
     "משתמש אחד יכול ליצור מתכונים רבים. "
     "כל מתכון שייך בדיוק למשתמש אחד (שדה UserId חובה). "
     "מחיקת משתמש מוחקת אוטומטית את כל המתכונים שלו (Cascade)."),
    ("Album → Recipe (1:N)",
     "אלבום אחד יכול להכיל מתכונים רבים. "
     "כל מתכון שייך לאלבום אחד בדיוק (שדה AlbumId חובה). "
     "מחיקת אלבום מוחקת אוטומטית את כל המתכונים שבו (Cascade)."),
    ("User ↔ Recipe (N:M דרך UserFavorites)",
     "משתמש יכול לסמן מתכונים רבים כמועדפים, "
     "ומתכון יכול להיות מסומן כמועדף על-ידי משתמשים רבים. "
     "הקשר ממומש דרך טבלת UserFavorites עם מפתח ראשי מורכב."),
    ("User → UserFavorite (1:N)",
     "משתמש אחד יכול להחזיק רשומות מועדפים רבות. "
     "מחיקת משתמש מוחקת את כל רשומות המועדפים שלו."),
    ("Recipe → UserFavorite (1:N)",
     "מתכון אחד יכול להופיע ברשומות מועדפים רבות. "
     "מחיקת מתכון מוחקת את כל הסימונים שלו כמועדף."),
]

for title, desc in relationships:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run(title + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

# ===========================
# 7.8 Enum RecipeType
# ===========================
add_subheading(doc, "7.8 ערכי Enum – RecipeType")

add_paragraph(doc,
    "שדה RecipeType בטבלת Recipes מוגדר בקוד C# כ-Enum, "
    "אך נשמר במסד הנתונים כטקסט (HasConversion<string>()) "
    "לשיפור קריאות הנתונים:"
)

enum_rows = [
    ("Link", "1", "מתכון מסוג קישור – מכיל URL לאתר חיצוני"),
    ("Text", "2", "מתכון מסוג טקסט – מכיל מצרכים והוראות הכנה מלאות"),
]
make_table(doc,
    ["ערך טקסט (DB)", "ערך מספרי (Enum)", "משמעות"],
    enum_rows
)

# ===========================
# 7.9 היסטוריית מיגרציות
# ===========================
add_subheading(doc, "7.9 היסטוריית מיגרציות")

add_paragraph(doc,
    "מסד הנתונים עבר שלוש מיגרציות מאז הקמתו:"
)

migrations = [
    ("20260306121636_InitialCreate",
     "יצירת ארבע הטבלאות הבסיסיות: Users, Albums, Recipes, UserFavorites. "
     "הגדרת כל המפתחות, האינדקסים והקשרים."),
    ("20260311113219_UpdateModel",
     "הוספת שדה ImagePath (text, nullable) לטבלת Albums."),
    ("20260311182831_AddUserName",
     "הוספת שדה UserName (text, nullable) לטבלת Users."),
]

for name, desc in migrations:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run(name + ": ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

doc.add_paragraph()

# ===========================
# 7.10 סיכום
# ===========================
add_subheading(doc, "7.10 סיכום מבנה הנתונים")

add_paragraph(doc,
    "מבנה הנתונים של RecipeBox מבוסס על ארבע טבלאות בלבד, "
    "המממשות מודל ישיר ויעיל. "
    "כל הקשרים מוגדרים עם מחיקה מדורגת (CASCADE) המבטיחה עקביות הנתונים. "
    "השימוש ב-UUID כמפתח ראשי למשתמשים (לעומת Integer לשאר הטבלאות) "
    "מספק אקראיות ומניעת חשיפת מספר הרשמות. "
    "ה-Enum RecipeType נשמר כטקסט לקריאות גבוהה, "
    "ואינדקסים ממוקדים מבטיחים ביצועים טובים בשאילתות הנפוצות."
)

doc.save("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")
print("Section 7 added successfully.")
