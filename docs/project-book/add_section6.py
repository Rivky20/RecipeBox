from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_subheading(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return p

def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        p.paragraph_format.right_to_left = True
    except:
        pass
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        p.paragraph_format.right_to_left = True
    except:
        pass
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

# Page break before section 6
doc.add_page_break()

# ===========================
# SECTION 6 MAIN HEADING
# ===========================
add_heading(doc, "6. פירוט על הפרויקט – איך המערכת עובדת", level=1)

add_paragraph(doc,
    "פרק זה מתאר את אופן הפעולה של מערכת RecipeBox על פי הקוד המפותח. "
    "ההסבר מבוסס על ניתוח מעמיק של הרכיבים, הבקרים, המודלים ושכבת השירותים, "
    "ומציג את זרימת הפעולה של המשתמש בצורה שיטתית ואקדמית."
)

# ===========================
# 6.1 תזרים כללי של המשתמש
# ===========================
add_subheading(doc, "6.1 תזרים כללי של המשתמש")

add_paragraph(doc,
    "כלל הפעולות במערכת מתחילות בדף הבית (HomePage). "
    "המשתמש יכול לגלוש במתכונים ללא הרשמה, אך פעולות כמו הוספת מתכון, "
    "סימון מועדפים ועריכה דורשות אימות מלא. "
    "להלן שלבי השימוש הכלליים:"
)

steps_general = [
    "המשתמש מגיע לדף הבית ורואה את רשימת המתכונים הקיימים במערכת.",
    "ניתן לחפש ולסנן מתכונים ללא כניסה לחשבון.",
    "לחיצה על מתכון פותחת את דף פרטי המתכון.",
    "לצורך הוספה, עריכה, מחיקה או סימון מועדפים – המשתמש מופנה להתחברות.",
    "לאחר התחברות, מופיעות פעולות נוספות בממשק בהתאם לתפקיד המשתמש.",
    "מנהל מערכת (Admin) ניגש לדשבורד הניהול דרך הניווט העליון.",
]
for s in steps_general:
    add_numbered(doc, s)

# ===========================
# 6.2 תהליך הרשמה והתחברות
# ===========================
add_subheading(doc, "6.2 תהליך הרשמה והתחברות")

add_paragraph(doc,
    "מערכת האימות מבוססת על JWT (JSON Web Tokens) עם תוקף של 7 ימים. "
    "הסיסמאות מוצפנות באמצעות BCrypt לפני שמירתן במסד הנתונים."
)

add_paragraph(doc, "שלבי ההרשמה:", bold_prefix="")
reg_steps = [
    "המשתמש ממלא טופס הרשמה: כתובת אימייל (חובה), סיסמה (חובה), שם משתמש (אופציונלי).",
    "הלקוח שולח בקשת POST לנתיב /api/auth/register.",
    "השרת בודק אם האימייל כבר קיים במסד הנתונים.",
    "אם האימייל חדש – הסיסמה עוברת גיבוב (hashing) עם BCrypt ונשמרת.",
    "נוצר רשומת משתמש חדשה עם תפקיד ברירת מחדל: 'User'.",
    "השרת מחזיר טוקן JWT חתום עם פרטי המשתמש (UserId, Email, Role).",
    "הלקוח שומר את הטוקן ב-Context ומשתמש בו לכל הבקשות המאומתות.",
]
for s in reg_steps:
    add_numbered(doc, s)

add_paragraph(doc, "שלבי ההתחברות:", bold_prefix="")
login_steps = [
    "המשתמש מזין אימייל וסיסמה בדף ההתחברות (LoginPage).",
    "הלקוח שולח בקשת POST לנתיב /api/auth/login.",
    "השרת מאתר את המשתמש לפי האימייל.",
    "אם המשתמש חסום (IsBlocked = true) – מוחזרת שגיאת 403 Forbidden.",
    "הסיסמה שהוזנה מושווית לגיבוב השמור באמצעות BCrypt.Verify.",
    "בהצלחה – מוחזר טוקן JWT חדש עם תוקף של 7 ימים.",
    "הלקוח מאחסן את הטוקן ומעדכן את ה-AuthContext לכל הרכיבים.",
]
for s in login_steps:
    add_numbered(doc, s)

# ===========================
# 6.3 תהליך יצירת מתכון
# ===========================
add_subheading(doc, "6.3 תהליך יצירת מתכון")

add_paragraph(doc,
    "הוספת מתכון מתבצעת דרך עמוד AddRecipePage. "
    "המערכת תומכת בשני סוגי מתכונים: קישור (Link) ומתכון מלא (Text)."
)

add_paragraph(doc, "שלבי יצירת מתכון:", bold_prefix="")
create_steps = [
    "המשתמש המחובר לוחץ על 'הוסף מתכון' בניווט.",
    "מוצג טופס הוספה עם שדות: שם, תיאור, סוג מתכון, תמונה.",
    "אם הסוג הוא 'קישור' – מתווסף שדה URL לאתר המתכון המקורי.",
    "אם הסוג הוא 'טקסט' – מתווספים שדות: מצרכים ורשימת הוראות הכנה.",
    "ניתן להעלות תמונה; הלקוח שולח אותה ל-/api/images/upload.",
    "השרת מעלה את התמונה ל-Cloudinary ומחזיר URL ציבורי.",
    "לאחר מילוי הטופס, נשלחת בקשת POST אל /api/recipes עם טוקן JWT בכותרת.",
    "השרת מאמת את הטוקן, שולף את מזהה המשתמש ושומר את המתכון עם UserId.",
    "המתכון נשמר במסד הנתונים עם חותמת זמן (CreatedAt) אוטומטית.",
    "המשתמש מופנה לדף הבית ומוצג המתכון החדש ברשימה.",
]
for s in create_steps:
    add_numbered(doc, s)

# ===========================
# 6.4 תהליך עריכת מתכון
# ===========================
add_subheading(doc, "6.4 תהליך עריכת מתכון")

add_paragraph(doc,
    "עריכת מתכון אפשרית רק לבעלים של המתכון או למנהל מערכת. "
    "הבדיקה מתבצעת בשרת לפני כל עדכון."
)

edit_steps = [
    "המשתמש לוחץ על כפתור 'ערוך' בדף פרטי המתכון.",
    "הטופס נטען עם הנתונים הקיימים של המתכון.",
    "המשתמש מבצע שינויים בשדות הרצויים.",
    "ניתן להחליף תמונה – מתבצעת העלאה חדשה ל-Cloudinary.",
    "לחיצה על 'שמור' שולחת בקשת PUT אל /api/recipes/{id} עם הטוקן.",
    "השרת בודק: האם UserId בטוקן תואם ל-UserId של המתכון, או שהתפקיד הוא Admin.",
    "אם לא מורשה – מוחזרת שגיאת 403 Forbidden.",
    "אם מורשה – הנתונים מתעדכנים במסד הנתונים.",
    "הלקוח מקבל אישור ומציג את המתכון המעודכן.",
]
for s in edit_steps:
    add_numbered(doc, s)

# ===========================
# 6.5 תהליך מחיקת מתכון
# ===========================
add_subheading(doc, "6.5 תהליך מחיקת מתכון")

add_paragraph(doc,
    "מחיקת מתכון כוללת מחיקה מדורגת (Cascade Delete) של כל הנתונים הקשורים, "
    "כולל רשומות המועדפים של המתכון."
)

delete_steps = [
    "המשתמש לוחץ על 'מחק' בדף פרטי המתכון.",
    "מוצגת תיבת דו-שיח לאישור המחיקה.",
    "לאחר אישור, נשלחת בקשת DELETE אל /api/recipes/{id} עם טוקן JWT.",
    "השרת מוודא שהמשתמש הוא בעלים של המתכון או מנהל מערכת.",
    "Entity Framework Core מוחק אוטומטית גם את רשומות UserFavorites הקשורות.",
    "מוחזרת תשובת 204 No Content לאישור המחיקה.",
    "הלקוח מסיר את המתכון מהרשימה המקומית.",
]
for s in delete_steps:
    add_numbered(doc, s)

# ===========================
# 6.6 חיפוש וסינון
# ===========================
add_subheading(doc, "6.6 חיפוש וסינון מתכונים")

add_paragraph(doc,
    "מנגנון החיפוש והסינון פועל בצד השרת. "
    "כל בקשת GET /api/recipes יכולה לקבל פרמטרים לחיפוש, סינון ומיון."
)

add_paragraph(doc, "אפשרויות החיפוש הזמינות:")
search_opts = [
    ("חיפוש טקסטואלי:", "חיפוש לפי שם מתכון או תיאור (חיפוש חלקי, case-insensitive)."),
    ("סינון לפי אלבום:", "ניתן לסנן את המתכונים לפי AlbumId ולראות רק מתכונים מאלבום מסוים."),
    ("מיון לפי שם:", "סידור המתכונים בסדר אלפביתי לפי שם."),
    ("מיון לפי תאריך:", "סידור לפי תאריך יצירה, החדש ביותר תחילה."),
    ("מיון חכם למשתמש מחובר:", "מתכונים של המשתמש עצמו מוצגים ראשונים, אחריהם מועדפים, ואז שאר המתכונים."),
]
for bold, text in search_opts:
    add_bullet(doc, text, bold_prefix=bold)

add_paragraph(doc,
    "הפרמטרים מועברים כ-Query String בכתובת הבקשה, "
    "ורכיב ה-SearchBar בלקוח מעדכן אותם בזמן אמת עם כל שינוי בשדה החיפוש."
)

# ===========================
# 6.7 הרשאות ובקרת גישה
# ===========================
add_subheading(doc, "6.7 הרשאות ובקרת גישה")

add_paragraph(doc,
    "המערכת מיישמת מודל הרשאות דו-שכבתי: "
    "אימות זהות (Authentication) באמצעות JWT, "
    "ואישור פעולה (Authorization) לפי תפקיד ובעלות."
)

add_paragraph(doc, "טבלת הרשאות לפי תפקיד:")

perms = [
    ("צפייה במתכונים", "כולם (גם ללא כניסה)", "כולם (גם ללא כניסה)"),
    ("חיפוש וסינון", "כולם", "כולם"),
    ("הוספת מתכון", "משתמש מחובר", "מנהל מערכת"),
    ("עריכת מתכון", "בעלים בלבד", "כל מתכון"),
    ("מחיקת מתכון", "בעלים בלבד", "כל מתכון"),
    ("ניהול מועדפים", "משתמש מחובר", "מנהל מערכת"),
    ("יצירת אלבום", "לא מורשה", "מנהל מערכת בלבד"),
    ("עריכת/מחיקת אלבום", "לא מורשה", "מנהל מערכת בלבד"),
    ("חסימת משתמשים", "לא מורשה", "מנהל מערכת בלבד"),
    ("צפייה בסטטיסטיקות", "לא מורשה", "מנהל מערכת בלבד"),
    ("העלאת תמונות", "משתמש מחובר", "מנהל מערכת"),
]

table = doc.add_table(rows=len(perms)+1, cols=3)
table.style = 'Table Grid'

hdr = table.rows[0].cells
hdr[0].text = "פעולה"
hdr[1].text = "משתמש רגיל"
hdr[2].text = "מנהל מערכת"
for cell in hdr:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True

for i, (action, user, admin) in enumerate(perms):
    row = table.rows[i+1].cells
    row[0].text = action
    row[1].text = user
    row[2].text = admin
    for cell in row:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

add_paragraph(doc,
    "מנגנון החסימה: משתמש שנחסם על ידי מנהל המערכת אינו יכול להתחבר גם אם הסיסמה נכונה. "
    "בקרת הגישה מיושמת גם בצד הלקוח (הסתרת כפתורים) וגם בצד השרת (בדיקה לפני כל פעולה)."
)

# ===========================
# 6.8 תזרים שימוש בעמודים
# ===========================
add_subheading(doc, "6.8 תזרים שימוש בעמודים")

add_paragraph(doc,
    "להלן סקירת העמודים המרכזיים ואופן פעולתם במערכת:"
)

pages = [
    (
        "דף הבית (HomePage)",
        "מציג את כל המתכונים הזמינים. "
        "כולל תיבת חיפוש, אפשרויות מיון וסינון לפי אלבום. "
        "כרטיסי המתכון מציגים תמונה, שם ותיאור קצר. "
        "לחיצה על כרטיס מנווטת לעמוד פרטי המתכון."
    ),
    (
        "עמוד פרטי מתכון",
        "מציג את כל מידע המתכון: שם, תיאור, מצרכים, הוראות הכנה. "
        "עבור מתכון מסוג 'קישור' – מוצג לינק לאתר החיצוני. "
        "כפתורי עריכה ומחיקה מוצגים רק לבעלים ולמנהל. "
        "כפתור 'הוסף למועדפים' מוצג למשתמשים מחוברים."
    ),
    (
        "עמוד הוספת מתכון (AddRecipePage)",
        "טופס דינמי שמשתנה לפי סוג המתכון שנבחר. "
        "כולל אפשרות העלאת תמונה, עם תצוגה מקדימה לפני השמירה. "
        "כולל כלי AI לשיפור הטקסט, הצעת שם ותיאור, ולהמרת יחידות מידה."
    ),
    (
        "עמוד המועדפים (FavoritesPage)",
        "מציג את כל המתכונים שהמשתמש סימן כמועדפים. "
        "ניתן להסיר מועדפים ישירות מהעמוד. "
        "הנתונים נשלפים מהנתיב /api/users/{userId}/favorites."
    ),
    (
        "עמוד ה-AI (AIPage)",
        "ממשק לחיפוש ויצירת מתכונים חכמים. "
        "המשתמש מזין רשימת מצרכים או שם מתכון. "
        "המערכת מחפשת בין המתכונים הקיימים תחילה. "
        "אם לא נמצא – נשלחת בקשה ל-Groq API (דרך פרוקסי) לקבלת מתכון מוגנרט."
    ),
    (
        "דשבורד ניהול (Admin)",
        "נגיש רק למשתמשים בעלי תפקיד Admin. "
        "כולל עמודים נפרדים לניהול משתמשים, אלבומים ומתכונים. "
        "מציג סטטיסטיקות מערכת: סה\"כ משתמשים, מתכונים, אלבומים ומשתמשים חסומים. "
        "פעולות: חסימה/שחרור משתמשים, יצירה/עריכה/מחיקה של אלבומים."
    ),
]

for title, desc in pages:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title + ": ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(desc)
    run2.font.size = Pt(11)

# ===========================
# 6.9 סיכום זרימת המערכת
# ===========================
add_subheading(doc, "6.9 סיכום")

add_paragraph(doc,
    "מערכת RecipeBox בנויה על עקרונות של הפרדת אחריות ברורה בין שכבות הלקוח והשרת. "
    "כל פעולה שדורשת הרשאה עוברת אימות כפול: "
    "ברמת הממשק (הסתרת אפשרויות) וברמת ה-API (בדיקת JWT ותפקיד). "
    "תכנון הנתונים תומך במחיקה מדורגת המבטיחה עקביות הנתונים בכל עת. "
    "שילוב יכולות ה-AI מאפשר למשתמשים לא רק לנהל מתכונים קיימים, "
    "אלא גם לגלות וליצור מתכונים חדשים בעזרת מודל שפה מתקדם."
)

doc.save("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")
print("Section 6 added successfully.")
