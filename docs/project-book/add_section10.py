from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_subheading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return p

def add_sub3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ===========================
# PAGE BREAK
# ===========================
doc.add_page_break()

# ===========================
# SECTION 10 HEADING
# ===========================
add_heading(doc, "10. מדריך משתמש לפי סוגי משתמשים", level=1)

add_paragraph(doc,
    "פרק זה מתאר את חוויית השימוש במערכת RecipeBox לפי כל סוג משתמש. "
    "המערכת מכירה בשלושה סוגי משתמשים: אורח (לא מחובר), משתמש רגיל ומנהל. "
    "כל סוג משתמש מקבל גישה שונה לתכונות המערכת, כפי שמוגדר בשרת דרך JWT ו-Role-Based Authorization."
)

# ===========================
# SUMMARY TABLE
# ===========================
add_subheading(doc, "10.1 טבלת סיכום הרשאות")

permissions_table = [
    ("צפייה במתכונים ואלבומים",     "✔", "✔", "✔"),
    ("חיפוש וסינון מתכונים",         "✔", "✔", "✔"),
    ("הרשמה / כניסה למערכת",         "✔", "—", "—"),
    ("יצירת מתכון חדש",              "✗", "✔", "✔"),
    ("עריכת מתכון עצמי",             "✗", "✔", "✔"),
    ("מחיקת מתכון עצמי",             "✗", "✔", "✔"),
    ("עריכת / מחיקת כל מתכון",       "✗", "✗", "✔"),
    ("הוספת מתכון למועדפים",         "✗", "✔", "✔"),
    ("צפייה ברשימת מועדפים",         "✗", "✔", "✔ (כל משתמש)"),
    ("העלאת תמונות",                 "✗", "✔", "✔"),
    ("שימוש ב-AI (עוזר מתכונים)",    "✗", "✔", "✔"),
    ("ניהול אלבומים (יצירה/עריכה/מחיקה)", "✗", "✗", "✔"),
    ("ניהול משתמשים (חסום/בטל חסימה)",   "✗", "✗", "✔"),
    ("צפייה בסטטיסטיקות מערכת",     "✗", "✗", "✔"),
    ("גישה לדפי ניהול (Admin Panel)", "✗", "✗", "✔"),
]

make_table(doc,
    ["פעולה", "אורח", "משתמש", "מנהל"],
    permissions_table
)

# ===========================
# 10.2 אורח
# ===========================
add_subheading(doc, "10.2 אורח (Unauthenticated)")

add_sub3(doc, "10.2.1 תיאור")
add_paragraph(doc,
    "אורח הוא כל גולש שמגיע לאתר מבלי להתחבר. "
    "הוא יכול לעיין בתוכן הציבורי של המערכת אך אינו יכול לבצע פעולות כתיבה. "
    "המערכת מזהה אותו כמשתמש לא מאומת — אין JWT בבקשותיו."
)

add_sub3(doc, "10.2.2 הרשאות")
add_bullet(doc, "קריאה של כל המתכונים והאלבומים (GET /api/recipes, GET /api/albums).")
add_bullet(doc, "חיפוש וסינון מתכונים לפי שם, תיאור, אלבום ומיון.")
add_bullet(doc, "צפייה בדף מתכון בודד (GET /api/recipes/{id}).")
add_bullet(doc, "הרשמה (POST /api/auth/register) וכניסה (POST /api/auth/login).")
add_bullet(doc, "אין גישה לפעולות הדורשות [Authorize]: יצירה, עריכה, מחיקה, מועדפים, AI.")

add_sub3(doc, "10.2.3 זרימת עבודה טיפוסית – אורח")
add_numbered(doc, "הגולש פותח את דף הבית (HomePage) — רואה את כל המתכונים.")
add_numbered(doc, "מחפש מתכון לפי מילת מפתח בשדה החיפוש.")
add_numbered(doc, "לוחץ על כרטיסיית מתכון ועובר לדף המתכון (RecipePage).")
add_numbered(doc, "מנסה ללחוץ על \"הוסף למועדפים\" — המערכת פותחת את חלון הכניסה (LoginModal).")
add_numbered(doc, "נרשם (RegisterPage) או מתחבר (LoginPage) ועובר להיות משתמש רגיל.")

# ===========================
# 10.3 משתמש רגיל
# ===========================
add_subheading(doc, "10.3 משתמש רגיל (Role: \"User\")")

add_sub3(doc, "10.3.1 תיאור")
add_paragraph(doc,
    "משתמש רגיל הוא כל מי שנרשם למערכת. "
    "לאחר הרשמה, השרת מקצה לו את התפקיד 'User' ושומר אותו בטוקן ה-JWT. "
    "הטוקן תקף ל-7 ימים ומאוחסן ב-localStorage בדפדפן. "
    "הוא יכול לנהל את המתכונים שלו בלבד ולא יכול לגשת לאיזורי ניהול."
)

add_sub3(doc, "10.3.2 הרשאות")
add_bullet(doc, "כל הרשאות האורח — צפייה, חיפוש, סינון.")
add_bullet(doc, "יצירת מתכון חדש — המתכון נוצר עם UserId של המשתמש המחובר.")
add_bullet(doc, "עריכת מתכונים שלו בלבד — ניסיון לערוך מתכון של אחר מחזיר 403 Forbidden.")
add_bullet(doc, "מחיקת מתכונים שלו בלבד — אותו כלל כמו עריכה.")
add_bullet(doc, "הוספה והסרה של מתכונים מרשימת המועדפים שלו.")
add_bullet(doc, "צפייה ברשימת המועדפים שלו (GET /api/users/{userId}/favorites).")
add_bullet(doc, "העלאת תמונות לשרת (POST /api/images/upload).")
add_bullet(doc, "שימוש בעוזר ה-AI לקבלת רעיונות ומתכונים (POST /api/ai/chat).")

add_sub3(doc, "10.3.3 זרימת יצירת מתכון")
add_numbered(doc, "המשתמש לוחץ על \"הוסף מתכון\" בניווט — עובר לדף AddRecipePage.")
add_numbered(doc, "בוחר סוג מתכון: טקסט (עם מרכיבים והוראות) או קישור (URL חיצוני).")
add_numbered(doc, "ממלא שם, תיאור ואלבום (חובה), ומעלה תמונה אופציונלית.")
add_numbered(doc, "לוחץ \"שמור\" — הלקוח שולח POST /api/recipes עם JWT.")
add_numbered(doc, "השרת מאמת את הטוקן, קורא את UserId מה-Claims, ויוצר את המתכון בשם המשתמש.")
add_numbered(doc, "המשתמש מועבר לדף המתכון החדש (RecipePage).")

add_sub3(doc, "10.3.4 זרימת עריכת מתכון")
add_numbered(doc, "המשתמש נמצא בדף מתכון שלו ורואה כפתור \"ערוך\" (מוצג רק לבעלים).")
add_numbered(doc, "לוחץ \"ערוך\" — עובר לדף EditRecipePage עם פרטי המתכון ממולאים.")
add_numbered(doc, "משנה את הפרטים הרצויים ולוחץ \"שמור\".")
add_numbered(doc, "הלקוח שולח PUT /api/recipes/{id} עם JWT.")
add_numbered(doc, "השרת בודק שה-userId ב-JWT תואם ל-UserId של המתכון, ומעדכן.")
add_numbered(doc, "המשתמש רואה את הנתונים המעודכנים בדף.")

add_sub3(doc, "10.3.5 זרימת ניהול מועדפים")
add_numbered(doc, "בכל כרטיסיית מתכון מופיע כפתור לב (FavoriteButton).")
add_numbered(doc, "לחיצה על הלב שולחת POST /api/users/{userId}/favorites עם recipeId.")
add_numbered(doc, "השרת מוודא שה-userId ב-URL שווה ל-userId ב-JWT (CanAccessUser).")
add_numbered(doc, "המתכון מתווסף לרשימה; הלב הופך למלא.")
add_numbered(doc, "ניתן לצפות בכל המועדפים בדף FavoritesPage.")
add_numbered(doc, "לחיצה שנייה על הלב מסירה את המתכון מהמועדפים (DELETE).")

# ===========================
# 10.4 מנהל
# ===========================
add_subheading(doc, "10.4 מנהל (Role: \"Admin\")")

add_sub3(doc, "10.4.1 תיאור")
add_paragraph(doc,
    "מנהל הוא משתמש שהוקצה לו תפקיד 'Admin' ישירות בבסיס הנתונים. "
    "אין דרך להפוך למנהל דרך ממשק המשתמש — הקצאת תפקיד זה נעשית ידנית. "
    "לאחר כניסה, הטוקן שלו מכיל Role=Admin ומעניק גישה לכל נקודות הקצה המוגנות. "
    "ממשק הניהול (Admin Panel) מוצג בניווט רק למנהלים."
)

add_sub3(doc, "10.4.2 הרשאות")
add_bullet(doc, "כל הרשאות המשתמש הרגיל.")
add_bullet(doc, "עריכה ומחיקה של כל מתכון — ללא קשר לבעלות.")
add_bullet(doc, "יצירת אלבומים חדשים (POST /api/albums) — מוגבל ל-Admin בלבד.")
add_bullet(doc, "עריכת פרטי אלבום: שם, תיאור, תמונת רקע (PUT /api/albums/{id}).")
add_bullet(doc, "מחיקת אלבום — גוררת מחיקת כל מתכוניו (Cascade Delete).")
add_bullet(doc, "צפייה ברשימת כל המשתמשים (GET /api/users).")
add_bullet(doc, "חסימת משתמש (PATCH /api/users/{id}/block) — משתמש חסום לא יכול להתחבר.")
add_bullet(doc, "ביטול חסימת משתמש (PATCH /api/users/{id}/unblock).")
add_bullet(doc, "צפייה בסטטיסטיקות מערכת (GET /api/users/stats).")
add_bullet(doc, "גישה למועדפים של כל משתמש (CanAccessUser מאפשר כניסה לכל userId).")

add_sub3(doc, "10.4.3 דפי ניהול")

admin_pages = [
    ("AdminStatsPage",   "/admin/stats",   "סטטיסטיקות מערכת: סה\"כ משתמשים, מתכונים, אלבומים, מועדפים"),
    ("AdminUsersPage",   "/admin/users",   "רשימת כל המשתמשים עם אפשרות חסימה/ביטול חסימה"),
    ("AdminAlbumsPage",  "/admin/albums",  "ניהול אלבומים: יצירה, עריכה, מחיקה, העלאת תמונת רקע"),
    ("AdminRecipesPage", "/admin/recipes", "צפייה ומחיקה של כל מתכוני המערכת"),
]
make_table(doc, ["דף", "נתיב", "תיאור"], admin_pages)

add_sub3(doc, "10.4.4 זרימת ניהול אלבומים")
add_numbered(doc, "המנהל עובר לדף AdminAlbumsPage — רואה טבלת כל האלבומים.")
add_numbered(doc, "לוחץ \"אלבום חדש\" — מופיע טופס עם שדות: שם (חובה), תיאור, תמונת רקע.")
add_numbered(doc, "יכול להדביק URL תמונה ישירות או להעלות קובץ תמונה מהמחשב.")
add_numbered(doc, "לוחץ \"שמור\" — נשלח POST /api/albums עם JWT מסוג Admin.")
add_numbered(doc, "האלבום החדש מופיע מיד בטבלה.")
add_numbered(doc, "לחיצה על עריכה (עט) מאכלסת את הטופס עם נתוני האלבום הקיים.")
add_numbered(doc, "לחיצה על מחיקה (פח) פותחת דיאלוג אישור לפני ביצוע הפעולה.")

add_sub3(doc, "10.4.5 זרימת ניהול משתמשים")
add_numbered(doc, "המנהל עובר לדף AdminUsersPage — רואה טבלה: אימייל, תפקיד, מספר מתכונים, סטטוס.")
add_numbered(doc, "משתמשים פעילים מסומנים בירוק, חסומים באדום.")
add_numbered(doc, "לחיצה על \"חסום\" — נשלח PATCH /api/users/{id}/block.")
add_numbered(doc, "המשתמש החסום לא יוכל להתחבר — ינוסה 401 'Your account has been blocked'.")
add_numbered(doc, "לחיצה על \"בטל חסימה\" — נשלח PATCH /api/users/{id}/unblock ומשחזר גישה.")
add_paragraph(doc,
    "שים לב: כפתורי חסימה אינם מוצגים ממשתמשים בעלי תפקיד Admin — "
    "ניתן לחסום רק משתמשים רגילים."
)

# ===========================
# 10.5 מצבי גבול
# ===========================
add_subheading(doc, "10.5 מצבי גבול וטיפול בשגיאות")

edge_cases = [
    ("משתמש מנסה לערוך מתכון של אחר",
     "השרת מחזיר 403 Forbidden; הלקוח מציג הודעת שגיאה."),
    ("משתמש חסום מנסה להתחבר",
     "השרת מחזיר 401 עם המסר 'Your account has been blocked'; ממשק מציג שגיאה."),
    ("טוקן JWT פג תוקף",
     "השרת מחזיר 401; Axios Interceptor מנקה את localStorage ומנתק את המשתמש."),
    ("מנהל מוחק אלבום עם מתכונים",
     "Cascade Delete מוחק את כל המתכונים בו ואת רשומות המועדפים הקשורות אליהם."),
    ("ניסיון הוספת מתכון כפול למועדפים",
     "השרת מחזיר 409 Conflict עם המסר 'Recipe is already in favorites'."),
    ("העלאת תמונה נכשלת",
     "Cloudinary מחזיר שגיאה; הלקוח מציג הודעה ומאפשר ניסיון חוזר."),
    ("אורח לוחץ על פעולה מוגנת",
     "ProtectedRoute מנתב לדף ראשי; LoginModal נפתח אוטומטית במקרים מסוימים."),
]
make_table(doc,
    ["תרחיש", "התנהגות המערכת"],
    edge_cases
)

# ===========================
# 10.6 סיכום
# ===========================
add_subheading(doc, "10.6 סיכום")

add_paragraph(doc,
    "מערכת RecipeBox מחלקת את המשתמשים לשלוש רמות הרשאה ברורות: "
    "אורח שיכול לצפות אך לא לפעול, משתמש רגיל שיכול לנהל את התוכן שלו, "
    "ומנהל שיכול לנהל את כל המערכת. "
    "ההרשאות נאכפות הן בצד השרת (דרך JWT, [Authorize] ובדיקות בעלות) "
    "והן בצד הלקוח (דרך ProtectedRoute והסתרת כפתורים לא רלוונטיים). "
    "הפרדה זו מבטיחה שכל משתמש רואה ויכול לבצע בדיוק את הפעולות המתאימות לרמתו."
)

doc.save("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")
print("Section 10 added successfully.")
