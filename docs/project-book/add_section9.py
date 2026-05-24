from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_subheading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    return p

def add_sub3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ===========================
# PAGE BREAK
# ===========================
doc.add_page_break()

# ===========================
# SECTION 9 HEADING
# ===========================
add_heading(doc, "9. מבנה המערכת", level=1)

add_paragraph(doc,
    "פרק זה מתאר את ארכיטקטורת המערכת הכוללת של RecipeBox — "
    "את רכיביה הראשיים, את אחריות כל שכבה, את מבנה התיקיות בפועל "
    "ואת זרימת התקשורת בין הלקוח, השרת ובסיס הנתונים. "
    "כל הפרטים מבוססים על קריאת קוד המקור בלבד."
)

# ===========================
# 9.1 סקירת הארכיטקטורה
# ===========================
add_subheading(doc, "9.1 סקירת הארכיטקטורה הכללית")

add_paragraph(doc,
    "המערכת בנויה בארכיטקטורת Client-Server קלאסית המחולקת לשלוש שכבות עיקריות:"
)
add_bullet(doc,
    "לקוח (Client) – אפליקציית React הרצה בדפדפן המשתמש. "
    "אחראית על כל ממשק המשתמש, ניהול המצב בצד הלקוח וקריאות ל-API.",
    bold_prefix="שכבת הלקוח:"
)
add_bullet(doc,
    "שרת (Server) – ASP.NET Core Web API הרץ ב-.NET 9. "
    "אחראי על לוגיקה עסקית, אימות, הרשאות וגישה לבסיס הנתונים.",
    bold_prefix="שכבת השרת:"
)
add_bullet(doc,
    "בסיס הנתונים (Database) – PostgreSQL. "
    "מאחסן את כל נתוני המערכת: משתמשים, אלבומים, מתכונים ומועדפים.",
    bold_prefix="שכבת הנתונים:"
)
add_paragraph(doc,
    "בנוסף, המערכת משתמשת בשני שירותים חיצוניים: "
    "Cloudinary לאחסון ותצוגת תמונות, ו-Groq AI לעיבוד שפה טבעית."
)

# ===========================
# 9.2 שכבת הלקוח (React)
# ===========================
add_subheading(doc, "9.2 שכבת הלקוח – Client (React)")

add_paragraph(doc,
    "הלקוח מפותח ב-React עם TypeScript ומשתמש ב-Chakra UI לעיצוב. "
    "הוא Single Page Application (SPA) המנוהל דרך React Router DOM."
)

add_sub3(doc, "9.2.1 מבנה תיקיות הלקוח")

add_paragraph(doc,
    "תיקיית client/src/ מאורגנת לפי אחריות פונקציונלית:"
)

client_structure = [
    ("client/src/pages/",            "דפי האפליקציה הראשיים (HomePage, RecipePage, AlbumPage, AIPage, FavoritesPage, Admin/*)"),
    ("client/src/components/",       "קומפוננטות ממשק משותפות, מחולקות לתת-תיקיות: layout/, recipes/, albums/, search/, common/"),
    ("client/src/components/layout/","Navbar, Footer, Layout, ProtectedRoute, LoginModal"),
    ("client/src/components/recipes/","RecipeCard, RecipeForm, FavoriteButton"),
    ("client/src/components/albums/","AlbumCard"),
    ("client/src/components/search/","SearchBar, SortSelect"),
    ("client/src/components/common/","ErrorMessage, Spinner, ConfirmDialog"),
    ("client/src/context/",          "AuthContext (מצב המשתמש המחובר), AuthModalContext (שליטה על חלון Login/Register)"),
    ("client/src/services/",         "קבצי גישה ל-API: axios.ts, authService.ts, recipeService.ts, albumService.ts, favoriteService.ts, userService.ts, imageService.ts, geminiService.ts"),
    ("client/src/types/",            "הגדרות טיפוסים TypeScript משותפות (index.ts)"),
    ("client/src/App.tsx",           "נקודת כניסה: הגדרת Routing, ChakraProvider, AuthProvider, AuthModalProvider"),
]
make_table(doc, ["נתיב", "תוכן ואחריות"], client_structure)

add_sub3(doc, "9.2.2 טכנולוגיות ואחריות – לקוח")

client_tech = [
    ("React 18",                "ספריית ממשק המשתמש; ניהול מחזור חיים קומפוננטות דרך Hooks"),
    ("TypeScript",              "הקלדה סטטית לכל קבצי הלקוח; מונע שגיאות טיפוס בזמן פיתוח"),
    ("React Router DOM",        "ניהול ניתוב בצד הלקוח (BrowserRouter, Routes, Route, Navigate, ProtectedRoute)"),
    ("Chakra UI v3",            "ספריית רכיבי UI עם תמיכה ב-RTL; משמשת לכל הטפסים, הכפתורים, ההתראות והדיאלוגים"),
    ("Axios",                   "ביצוע בקשות HTTP ל-API; כולל Interceptors לצירוף JWT וטיפול ב-401"),
    ("Context API",             "ניהול מצב גלובלי (אימות, מצב modal) ללא Redux"),
    ("Vite",                    "כלי Build ו-Dev Server – מחליף Create React App"),
    ("localStorage",            "שמירת מצב המשתמש המחובר (token, userId, role, email, userName) בין רענוני דפדפן"),
]
make_table(doc, ["טכנולוגיה", "אחריות"], client_tech)

# ===========================
# 9.3 שכבת השרת (ASP.NET Core)
# ===========================
add_subheading(doc, "9.3 שכבת השרת – Server (ASP.NET Core Web API)")

add_paragraph(doc,
    "השרת מפותח ב-C# ומבוסס על ASP.NET Core 9. "
    "הוא מיישם ארכיטקטורת שכבות (Layered Architecture) "
    "עם הפרדה בין Controllers, Services ו-Models."
)

add_sub3(doc, "9.3.1 מבנה תיקיות השרת")

server_structure = [
    ("server/RecipeBoxServer/Controllers/", "בקרים: AuthController, RecipesController, AlbumsController, FavoritesController, UsersController, ImagesController, AIController"),
    ("server/RecipeBoxServer/Services/",    "לוגיקה עסקית: AuthService, RecipeService, AlbumService, FavoriteService, UserService, CloudinaryService + ממשקי IService"),
    ("server/RecipeBoxServer/Models/",      "מודלי EF Core: User, Recipe, Album, UserFavorite, RecipeType (Enum)"),
    ("server/RecipeBoxServer/DTOs/",        "אובייקטי העברת נתונים: Auth/, Recipes/, Albums/, Favorites/, Users/, Admin/"),
    ("server/RecipeBoxServer/Data/",        "RecipeBoxDbContext – הגדרת DbSets, מפתחות וקשרים"),
    ("server/RecipeBoxServer/Migrations/",  "מיגרציות EF Core (InitialCreate, UpdateModel, AddUserName)"),
    ("server/RecipeBoxServer/Program.cs",   "נקודת כניסה: רישום שירותים, הגדרת JWT, CORS, Swagger והפעלת Auto-migrate"),
]
make_table(doc, ["נתיב", "תוכן ואחריות"], server_structure)

add_sub3(doc, "9.3.2 טכנולוגיות ואחריות – שרת")

server_tech = [
    ("ASP.NET Core 9",                 "פלטפורמת ה-Web API; pipeline: CORS → HTTPS → Authentication → Authorization → Controllers"),
    ("Entity Framework Core",          "ORM לגישה ל-PostgreSQL; מנהל מיגרציות, ניווט בין ישויות ו-LINQ queries"),
    ("Npgsql EF Core Provider",        "ספק ה-EF Core ל-PostgreSQL"),
    ("BCrypt.Net-Next",                "גיבוב (Hashing) וכן אימות סיסמאות בעת רישום וכניסה"),
    ("Microsoft.AspNetCore.Authentication.JwtBearer", "אימות JWT; אימות חתימה (HMAC-SHA256) ותוקף הטוקן"),
    ("Cloudinary SDK",                 "CloudinaryService: העלאת קבצי תמונה ל-CDN והחזרת URL ציבורי"),
    ("HttpClient",                     "AIController: בקשות HTTP פרוקסי לשירות Groq AI"),
    ("Swashbuckle / Swagger",          "תיעוד API אינטראקטיבי זמין בפיתוח בכתובת /"),
    ("Dependency Injection",           "כל השירותים רשומים כ-Scoped ב-Program.cs ומוזרקים לבקרים דרך קונסטרוקטור"),
]
make_table(doc, ["טכנולוגיה", "אחריות"], server_tech)

# ===========================
# 9.4 שכבת בסיס הנתונים
# ===========================
add_subheading(doc, "9.4 שכבת בסיס הנתונים – Database (PostgreSQL)")

add_paragraph(doc,
    "בסיס הנתונים הוא PostgreSQL, המנוהל לחלוטין דרך Entity Framework Core. "
    "המבנה מוגדר ב-RecipeBoxDbContext ומכיל ארבעה טבלאות עיקריות."
)

add_sub3(doc, "9.4.1 טבלאות ושדות מרכזיים")

db_tables = [
    ("Users",        "Id (Guid, PK), Email (unique), UserName, PasswordHash, Role ('User'/'Admin'), IsBlocked (bool), CreatedAt"),
    ("Albums",       "Id (int, PK), Name, Description, ImagePath"),
    ("Recipes",      "Id (int, PK), Name, Description, ImagePath, RecipeType (string: 'Text'/'Link'), Link, Ingredients, Instructions, AlbumId (FK), UserId (FK), CreatedAt"),
    ("UserFavorites","UserId (FK, PK חלקי), RecipeId (FK, PK חלקי) — מפתח ראשי מורכב"),
]
make_table(doc, ["טבלה", "שדות עיקריים"], db_tables)

add_sub3(doc, "9.4.2 קשרים בין טבלאות")

add_bullet(doc, "Album → Recipes: יחס 1:N. מחיקת אלבום גוררת מחיקת מתכוניו (Cascade).")
add_bullet(doc, "User → Recipes: יחס 1:N. מחיקת משתמש גוררת מחיקת מתכוניו (Cascade).")
add_bullet(doc, "User → UserFavorites: יחס 1:N עם Cascade.")
add_bullet(doc, "Recipe → UserFavorites: יחס 1:N עם Cascade.")
add_bullet(doc, "UserFavorites: מפתח ראשי מורכב (UserId, RecipeId) — מבטיח ייחודיות מועדף לכל משתמש.")
add_bullet(doc, "RecipeType שמור כמחרוזת ב-DB (HasConversion<string>()) — ערכים: 'Text' או 'Link'.")
add_bullet(doc, "שדה Email בטבלת Users מסומן עם Index ייחודי (IsUnique).")

# ===========================
# 9.5 זרימת תקשורת
# ===========================
add_subheading(doc, "9.5 זרימת התקשורת")

add_paragraph(doc,
    "כל התקשורת בין הלקוח לשרת מתבצעת דרך HTTP/HTTPS בפורמט JSON. "
    "אין שימוש ב-WebSockets, Server-Sent Events או גישה ישירה לבסיס הנתונים מהלקוח."
)

add_sub3(doc, "9.5.1 זרימת בקשה רגילה (ללא אימות)")

add_paragraph(doc, "דוגמה: טעינת רשימת מתכונים בדף הבית.")
add_bullet(doc, "הלקוח (HomePage) קורא ל-recipeService.getAll() עם פרמטרי סינון ומיון.")
add_bullet(doc, "recipeService שולח GET /api/recipes?... דרך מופע Axios המוגדר ב-axios.ts.")
add_bullet(doc, "Axios Request Interceptor מוסיף Authorization: Bearer <token> אם קיים token ב-localStorage.")
add_bullet(doc, "השרת מעבד: CORS → UseAuthentication (קריאת JWT, אופציונלי) → RecipesController.GetAll().")
add_bullet(doc, "RecipesController קורא ל-RecipeService.GetAllAsync() שמבצעת שאילתת LINQ לבסיס הנתונים.")
add_bullet(doc, "EF Core מתרגמת את ה-LINQ ל-SQL ומחזירה תוצאות מ-PostgreSQL.")
add_bullet(doc, "השרת מחזיר רשימת RecipeDto בפורמט JSON עם קוד 200 OK.")
add_bullet(doc, "הלקוח מקבל את הנתונים ומציג אותם דרך קומפוננטות RecipeCard.")

add_sub3(doc, "9.5.2 זרימת בקשה מאומתת (עם JWT)")

add_paragraph(doc, "דוגמה: הוספת מתכון למועדפים.")
add_bullet(doc, "המשתמש לוחץ על FavoriteButton; הקומפוננט קוראת ל-favoriteService.add(userId, recipeId).")
add_bullet(doc, "Axios מצרף את ה-JWT מ-localStorage כ-Authorization header.")
add_bullet(doc, "FavoritesController מקבל את הבקשה; UseAuthentication מפענחת את ה-JWT ומאכלסת את User.Claims.")
add_bullet(doc, "CanAccessUser(userId) מוודאת שה-Claim תואם את userId ב-URL.")
add_bullet(doc, "FavoriteService.AddFavoriteAsync() בודקת כפילות ויוצרת רשומת UserFavorite.")
add_bullet(doc, "EF Core מבצעת INSERT ל-PostgreSQL; השרת מחזיר 200 OK.")
add_bullet(doc, "Axios Response Interceptor: אם השרת מחזיר 401 – מנקה localStorage ושולח אירוע 'session-expired'; AuthContext מנתק את המשתמש.")

add_sub3(doc, "9.5.3 זרימת העלאת תמונה")

add_bullet(doc, "הלקוח שולח POST /api/images/upload עם קובץ בפורמט multipart/form-data.")
add_bullet(doc, "ImagesController מקבל את הקובץ ומעביר אותו ל-CloudinaryService.")
add_bullet(doc, "CloudinaryService מעלה את הקובץ ל-CDN של Cloudinary ומחזיר URL ציבורי.")
add_bullet(doc, "ה-URL מוחזר ללקוח ונשמר ב-imagePath של המתכון.")

add_sub3(doc, "9.5.4 זרימת גישה ל-AI")

add_bullet(doc, "הלקוח (geminiService.ts) שולח POST /api/ai/chat עם הודעת המשתמש.")
add_bullet(doc, "AIController מקבל את הגוף כ-JsonElement ומעביר אותו כפרוקסי ל-Groq AI.")
add_bullet(doc, "Groq מחזיר תגובת JSON; AIController מחלץ את content ומחזיר אותו ללקוח.")
add_bullet(doc, "geminiService.ts מנקה markdown עטיפה (```json```) ומבצע JSON.parse על התוצאה.")

# ===========================
# 9.6 אחריות טכנולוגיות
# ===========================
add_subheading(doc, "9.6 טבלת אחריות טכנולוגית")

responsibilities = [
    ("ניתוב ודפים",        "React Router DOM",              "client/src/App.tsx"),
    ("ממשק משתמש",         "React + Chakra UI",             "client/src/components/, pages/"),
    ("ניהול מצב גלובלי",   "React Context API",             "client/src/context/"),
    ("בקשות HTTP",          "Axios + Interceptors",          "client/src/services/axios.ts"),
    ("גישה ל-API (לקוח)",  "Service Layer (TypeScript)",    "client/src/services/*.ts"),
    ("הרשאות (לקוח)",      "ProtectedRoute",                "client/src/components/layout/ProtectedRoute.tsx"),
    ("שמירת Session",       "localStorage",                  "client/src/context/AuthContext.tsx"),
    ("ניתוב HTTP (שרת)",   "ASP.NET Core Controllers",      "server/.../Controllers/"),
    ("לוגיקה עסקית",       "Service Classes (C#)",          "server/.../Services/"),
    ("ORM ושאילתות DB",     "Entity Framework Core",         "server/.../Data/RecipeBoxDbContext.cs"),
    ("אימות זהות",          "JWT Bearer + BCrypt",           "server/.../Services/AuthService.cs"),
    ("הרשאות (שרת)",       "[Authorize] + Ownership Check", "server/.../Controllers/ + Services/"),
    ("בסיס נתונים",         "PostgreSQL",                    "חיצוני – חיבור דרך Connection String"),
    ("אחסון תמונות",        "Cloudinary CDN",                "server/.../Services/CloudinaryService.cs"),
    ("בינה מלאכותית",       "Groq AI (Proxy)",               "server/.../Controllers/AIController.cs"),
    ("תיעוד API",           "Swagger / OpenAPI",             "זמין בפיתוח בכתובת /"),
]
make_table(doc,
    ["אחריות", "טכנולוגיה", "מיקום בקוד"],
    responsibilities
)

# ===========================
# 9.7 סיכום
# ===========================
add_subheading(doc, "9.7 סיכום")

add_paragraph(doc,
    "ארכיטקטורת RecipeBox מבוססת על הפרדה ברורה בין שלוש שכבות: "
    "לקוח React האחראי לחוויית המשתמש וניהול מצב מקומי, "
    "שרת ASP.NET Core האחראי ללוגיקה עסקית, אימות והרשאות, "
    "ובסיס נתונים PostgreSQL המנוהל לחלוטין דרך EF Core. "
    "התקשורת בין השכבות מתבצעת אך ורק דרך REST API ב-JSON, "
    "כאשר JWT מאבטח כל בקשה מאומתת. "
    "שירותים חיצוניים (Cloudinary, Groq AI) משולבים דרך שכבת השרת בלבד, "
    "כך שמפתחות ה-API אינם חשופים ללקוח."
)

doc.save("docs/project-book/RECIPEBOX_PROJECT_BOOK.docx")
print("Section 9 added successfully.")
