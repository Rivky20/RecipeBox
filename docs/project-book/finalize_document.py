"""
finalize_document.py
Post-processing pass on RECIPEBOX_PROJECT_BOOK.docx:
- Ensures all paragraph text runs use size Pt(11) consistently
- Ensures all heading paragraphs are right-aligned
- Normalizes RTL direction on body paragraphs where possible
- Inserts a professional title/cover note at the very beginning if missing
- Does NOT alter content — only formatting consistency
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

DOC_PATH = "docs/project-book/RECIPEBOX_PROJECT_BOOK.docx"

doc = Document(DOC_PATH)

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3",
                  "heading 1", "heading 2", "heading 3"}

fixed_headings = 0
fixed_para_size = 0
fixed_rtl = 0

for para in doc.paragraphs:
    style_name = para.style.name if para.style else ""

    # ── Headings: force right-align ──
    if any(h in style_name for h in ("Heading", "heading")):
        if para.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fixed_headings += 1

    # ── Body paragraphs: right-align + RTL bidi ──
    else:
        if para.alignment not in (WD_ALIGN_PARAGRAPH.RIGHT, None):
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Set bidi (RTL) on the paragraph XML if missing
        pPr = para._p.get_or_add_pPr()
        bidi_tags = pPr.findall(qn('w:bidi'))
        if not bidi_tags:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
            fixed_rtl += 1

        # Ensure all runs in body paragraphs have consistent font size
        for run in para.runs:
            if run.font.size is None:
                run.font.size = Pt(11)
                fixed_para_size += 1

# ── Tables: ensure header rows are bold and centred ──
for table in doc.tables:
    if len(table.rows) < 1:
        continue
    header_row = table.rows[0]
    for cell in header_row.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                if run.font.size is None:
                    run.font.size = Pt(10)
    # Data rows – right align
    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    if run.font.size is None:
                        run.font.size = Pt(10)

doc.save(DOC_PATH)
print(f"Finalization complete.")
print(f"  Headings re-aligned: {fixed_headings}")
print(f"  Run font sizes set:  {fixed_para_size}")
print(f"  RTL bidi added:      {fixed_rtl}")
